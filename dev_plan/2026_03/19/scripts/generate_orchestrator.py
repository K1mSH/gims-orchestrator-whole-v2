# -*- coding: utf-8 -*-
"""
sync-orchestrator NGIS D23 클래스명세서 전체 생성 스크립트
- 7개 도메인 그룹, 43개 클래스
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

YELLOW_SHADING = 'FFFF00'
COL_WIDTHS_DXA = [1800, 2701, 1529, 2902]
ROW_HEIGHTS_WITH_OPS = [311, 321, 297, 297, 1300, 261, 275, 261, 258]
ROW_HEIGHTS_NO_OPS = [311, 321, 297, 297, 1300, 261, 275]


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _make_run(p, text, bold=False, font_size=9):
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
    rPr.append(rFonts)
    return run


def set_cell_text(cell, text, bold=False, font_size=9, alignment=None):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _make_run(p, text, bold=bold, font_size=font_size)


def set_cell_multiline(cell, lines_list, bold=False, font_size=9):
    for p in cell.paragraphs:
        p.clear()
    for i, line in enumerate(lines_list):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _make_run(p, line, bold=bold, font_size=font_size)


def merge_cells_horizontal(table, row_idx, start_col, end_col):
    table.cell(row_idx, start_col).merge(table.cell(row_idx, end_col))


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))


def set_table_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing = tblPr.find(qn('w:tblLayout'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(layout)


def set_table_indent(table, indent_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    indent_xml = f'<w:tblInd {nsdecls("w")} w:w="{indent_dxa}" w:type="dxa"/>'
    existing = tblPr.find(qn('w:tblInd'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(indent_xml))


def set_row_height(row, height_dxa):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{height_dxa}" w:hRule="atLeast"/>')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def set_column_widths(table, widths_dxa):
    tbl = table._tbl
    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in widths_dxa:
        grid_xml += f'<w:gridCol w:w="{w}"/>'
    grid_xml += '</w:tblGrid>'
    tblPr = tbl.tblPr
    tblPr.addnext(parse_xml(grid_xml))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_dxa):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                tcPr.append(parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{widths_dxa[i]}" w:type="dxa"/>'
                ))


def set_cell_border_nil(cell, sides):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_parts = []
    for side in sides:
        borders_parts.append(f'<w:{side} w:val="nil"/>')
    borders_xml = f'<w:tcBorders {nsdecls("w")}>{"".join(borders_parts)}</w:tcBorders>'
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(borders_xml))


def create_class_table(doc, package_name, class_id, class_name, screen_id,
                       class_type, class_overview, attributes_text,
                       operations_text=None):
    """Create a single class specification table."""
    has_ops = operations_text is not None
    num_rows = 9 if has_ops else 7

    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table)
    set_table_layout_fixed(table)
    set_table_indent(table, 125)
    set_column_widths(table, COL_WIDTHS_DXA)

    heights = ROW_HEIGHTS_WITH_OPS if has_ops else ROW_HEIGHTS_NO_OPS
    for i, row in enumerate(table.rows):
        if i < len(heights):
            set_row_height(row, heights[i])

    # Row 0: Header
    merge_cells_horizontal(table, 0, 0, 3)
    cell = table.cell(0, 0)
    set_cell_text(cell, '클래스 명세서[설계]', bold=True, font_size=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 1: 패키지명 / 클래스ID
    set_cell_text(table.cell(1, 0), '패키지명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(1, 0), YELLOW_SHADING)
    set_cell_text(table.cell(1, 1), package_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 2), '클래스ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(1, 2), YELLOW_SHADING)
    set_cell_text(table.cell(1, 3), class_id, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 2: 클래스명 / 화면ID
    set_cell_text(table.cell(2, 0), '클래스명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(2, 0), YELLOW_SHADING)
    set_cell_text(table.cell(2, 1), class_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 2), '화면ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(2, 2), YELLOW_SHADING)
    set_cell_text(table.cell(2, 3), screen_id, font_size=9)

    # Row 3: 클래스 타입
    set_cell_text(table.cell(3, 0), '클래스 타입', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(3, 0), YELLOW_SHADING)
    set_cell_text(table.cell(3, 1), class_type, font_size=9)
    set_cell_text(table.cell(3, 2), '', font_size=9)
    set_cell_text(table.cell(3, 3), '', font_size=9)
    set_cell_border_nil(table.cell(3, 1), ['right'])
    set_cell_border_nil(table.cell(3, 2), ['left', 'right'])
    set_cell_border_nil(table.cell(3, 3), ['left'])

    # Row 4: 클래스 개요
    cell0 = table.cell(4, 0)
    set_cell_shading(cell0, YELLOW_SHADING)
    set_cell_multiline(cell0, ['', '', '클래스 개요'], bold=True, font_size=9)
    for p in cell0.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    merge_cells_horizontal(table, 4, 1, 3)
    overview_lines = class_overview.split('\n')
    set_cell_multiline(table.cell(4, 1), overview_lines, font_size=9)

    # Row 5: 속성(Attributes) header
    merge_cells_horizontal(table, 5, 0, 3)
    cell = table.cell(5, 0)
    set_cell_text(cell, '속성(Attributes)', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 6: attributes content
    merge_cells_horizontal(table, 6, 0, 3)
    attr_lines = attributes_text.split('\n')
    set_cell_multiline(table.cell(6, 0), attr_lines, font_size=9)

    if has_ops:
        # Row 7: 오퍼레이션(Operations) header
        merge_cells_horizontal(table, 7, 0, 3)
        cell = table.cell(7, 0)
        set_cell_text(cell, '오퍼레이션(Operations)', bold=True, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, YELLOW_SHADING)

        # Row 8: operations content
        merge_cells_horizontal(table, 8, 0, 3)
        ops_lines = operations_text.split('\n')
        set_cell_multiline(table.cell(8, 0), ops_lines, font_size=9)

    doc.add_paragraph()  # spacing


# ============================================================
# Overview helper: 개요 + 상세내용 포맷
# ============================================================
def make_overview(line1, line2):
    return '\n'.join([
        '개요',
        f'-    {line1}',
        '',
        '상세내용',
        f'-    {line2}',
    ])


# ============================================================
# DATA DEFINITIONS
# ============================================================

def build_all_classes():
    """Return list of (group_name, class_type_label, class_data_list) tuples."""
    groups = []

    # ================================================================
    # GROUP 1: Agent
    # ================================================================
    agent_classes = []

    # --- AgentController ---
    agent_classes.append(('Controller', [{
        'class_name': 'AgentController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'Agent Controller를 정의한다.',
            'Agent CRUD, 헬스체크, Retention 설정, select-tables 조회, 테스트 데이터 생성/정리 기능을 제공한다.'
        ),
        'attributes': 'private final AgentService agentService;',
        'operations': '\n'.join([
            '// Agent 전체 목록 조회',
            '@GetMapping',
            'public ResponseEntity<List<AgentDto.Response>> getAgents();',
            '',
            '// Agent 단건 조회',
            '@GetMapping("/{id}")',
            'public ResponseEntity<AgentDto.Response> getAgent(@PathVariable Long id);',
            '',
            '// Agent 프로세스 agentCode 목록 조회',
            '@GetMapping("/discover")',
            'public ResponseEntity<Map<String, Object>> discoverAgents(@RequestParam String endpointUrl);',
            '',
            '// Agent 생성',
            '@PostMapping',
            'public ResponseEntity<AgentDto.Response> createAgent(@Valid @RequestBody AgentDto.CreateRequest request);',
            '',
            '// Agent 수정',
            '@PutMapping("/{id}")',
            'public ResponseEntity<AgentDto.Response> updateAgent(@PathVariable Long id, @RequestBody AgentDto.UpdateRequest request);',
            '',
            '// Agent 삭제',
            '@DeleteMapping("/{id}")',
            'public ResponseEntity<Void> deleteAgent(@PathVariable Long id);',
            '',
            '// Agent 헬스체크',
            '@PostMapping("/{id}/health-check")',
            'public ResponseEntity<AgentDto.HealthCheckResponse> healthCheck(@PathVariable Long id);',
            '',
            '// Retention 설정 조회',
            '@GetMapping("/{id}/retention")',
            'public ResponseEntity<?> getRetentionConfig(@PathVariable Long id);',
            '',
            '// Retention 설정 저장',
            '@PutMapping("/{id}/retention")',
            'public ResponseEntity<?> updateRetentionConfig(@PathVariable Long id, @RequestBody String body);',
            '',
            '// select-tables 조회',
            '@GetMapping("/{id}/select-tables")',
            'public ResponseEntity<?> getSelectTables(@PathVariable Long id);',
            '',
            '// 테스트 데이터 생성',
            '@PostMapping("/{id}/generate-test-data")',
            'public ResponseEntity<?> generateTestData(@PathVariable Long id, @RequestParam(defaultValue="1000") int count);',
            '',
            '// 테스트 데이터 정리',
            '@DeleteMapping("/{id}/clear-test-data")',
            'public ResponseEntity<?> clearTestData(@PathVariable Long id);',
        ]),
    }]))

    # --- AgentService ---
    agent_classes.append(('Service', [{
        'class_name': 'AgentService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'Agent 서비스를 정의한다.',
            'Agent CRUD, 헬스체크, Retention 설정 관리, 테스트 데이터 생성/정리, select-tables 조회 비즈니스 로직을 처리한다.'
        ),
        'attributes': '\n'.join([
            'private final AgentRepository agentRepository;',
            'private final DatasourceRepository datasourceRepository;',
            'private final DatasourceTableRepository tableRepository;',
            'private final PasswordEncryptor passwordEncryptor;',
            'private final RestTemplate restTemplate;',
            'private final EntityManager entityManager;',
            'private final ObjectMapper objectMapper;',
        ]),
        'operations': '\n'.join([
            '// 전체 Agent 목록 조회',
            'public List<AgentDto.Response> findAll();',
            '',
            '// 온라인 Agent 목록 조회',
            'public List<AgentDto.Response> findOnlineAgents();',
            '',
            '// ID로 Agent 단건 조회',
            'public AgentDto.Response findById(Long id);',
            '',
            '// agentCode로 Agent 조회',
            'public AgentDto.Response findByAgentCode(String agentCode);',
            '',
            '// Agent 생성',
            'public AgentDto.Response create(AgentDto.CreateRequest request);',
            '',
            '// Agent 프로세스 agentCode 목록 조회',
            'public Map<String, Object> discoverAgents(String endpointUrl);',
            '',
            '// Agent 수정',
            'public AgentDto.Response update(Long id, AgentDto.UpdateRequest request);',
            '',
            '// Agent 삭제',
            'public void delete(Long id);',
            '',
            '// Agent 헬스체크',
            'public AgentDto.HealthCheckResponse healthCheck(Long id);',
            '',
            '// Agent 상태 업데이트',
            'public void updateStatus(Long id, AgentStatus status);',
            '',
            '// Retention 설정 조회',
            'public Map<String, Object> getRetentionConfig(Long id);',
            '',
            '// Retention 설정 저장',
            'public Map<String, Object> updateRetentionConfig(Long id, String configJson);',
            '',
            '// 테스트 데이터 생성',
            'public Map<String, Object> generateTestData(Long id, int count);',
            '',
            '// 테스트 데이터 정리',
            'public Map<String, Object> clearTestData(Long id);',
            '',
            '// select-tables 조회',
            'public List<DatasourceDto.TableResponse> getSelectTables(Long agentId);',
        ]),
    }]))

    # --- AgentRequest ---
    agent_classes.append(('Request', [{
        'class_name': 'AgentRequest',
        'class_type': 'Class <<Request>>',
        'overview': make_overview(
            'Agent Request 정보를 정의한다.',
            'Agent 생성/수정 요청에 필요한 코드, 이름, URL, 존, 유형, 데이터소스, 테이블 ID 등의 파라미터를 정의한다.'
        ),
        'attributes': '\n'.join([
            'private String agentCode;           // 에이전트 고유 코드',
            'private String agentName;           // 에이전트 표시명',
            'private String endpointUrl;         // REST 엔드포인트 URL',
            'private String zone;                // 네트워크 존',
            'private AgentType agentType;        // 에이전트 유형',
            'private String datasourceTag;       // 데이터소스 태그',
            'private Boolean isActive;           // 활성화 여부',
            'private String sourceDatasourceId;  // 소스 데이터소스 ID',
            'private String targetDatasourceId;  // 타겟 데이터소스 ID',
            'private String description;         // 설명',
            'private AgentStatus status;         // 상태',
            'private List<Long> sourceTableIds;  // 소스 테이블 ID 목록',
            'private List<Long> targetTableIds;  // 타겟 테이블 ID 목록',
        ]),
        'operations': None,
    }]))

    # --- AgentResponse ---
    agent_classes.append(('Response', [{
        'class_name': 'AgentResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'Agent Response 정보를 정의한다.',
            'Agent 조회 응답에 포함되는 ID, 코드, 이름, 상태, 실행이력, 헬스체크 결과 등을 정의한다.'
        ),
        'attributes': '\n'.join([
            'private Long id;                         // PK',
            'private String agentCode;                // 에이전트 고유 코드',
            'private String agentName;                // 에이전트 표시명',
            'private String endpointUrl;              // REST 엔드포인트 URL',
            'private String zone;                     // 네트워크 존',
            'private Boolean isActive;                // 활성화 여부',
            'private AgentType agentType;             // 에이전트 유형',
            'private String datasourceTag;            // 데이터소스 태그',
            'private String sourceDatasourceId;       // 소스 데이터소스 ID',
            'private String targetDatasourceId;       // 타겟 데이터소스 ID',
            'private String description;              // 설명',
            'private AgentStatus status;              // 상태',
            'private LocalDateTime lastExecutedAt;    // 마지막 실행 시각',
            'private String lastExecutionStatus;      // 마지막 실행 결과',
            'private LocalDateTime createdAt;         // 생성 시각',
            'private List<Long> sourceTableIds;       // 소스 테이블 ID 목록',
            'private List<Long> targetTableIds;       // 타겟 테이블 ID 목록',
            'private String retentionConfig;          // Retention 설정 JSON',
            'private String message;                  // 헬스체크 결과 메시지',
        ]),
        'operations': None,
    }]))

    # --- Agent Entity ---
    agent_classes.append(('Entity', [
        {
            'class_name': 'Agent',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'Agent 엔티티를 정의한다.',
                '테이블: agent. 동기화 에이전트 정보를 관리하는 JPA 엔티티로, 코드/이름/URL/존/유형/상태/실행이력/Retention 설정을 포함한다. Schedule, AgentChainMember, AgentTable과 1:N 연관관계를 가진다.'
            ),
            'attributes': '\n'.join([
                '@Id @GeneratedValue(strategy = GenerationType.IDENTITY)',
                'private Long id;                              // 에이전트 PK',
                '',
                '@Column(name = "agent_code", unique = true, nullable = false)',
                'private String agentCode;                     // 에이전트 고유 코드',
                '',
                '@Column(name = "agent_name", nullable = false)',
                'private String agentName;                     // 에이전트 표시명',
                '',
                '@Column(name = "endpoint_url", nullable = false)',
                'private String endpointUrl;                   // REST 엔드포인트 URL',
                '',
                '@Column(name = "zone", nullable = false)',
                'private String zone;                          // 네트워크 존 (DMZ/INTERNAL)',
                '',
                '@Column(name = "is_active")',
                'private Boolean isActive;                     // 활성화 여부',
                '',
                '@Enumerated(EnumType.STRING)',
                '@Column(name = "agent_type", nullable = false)',
                'private AgentType agentType;                  // 에이전트 유형',
                '',
                '@Column(name = "datasource_tag")',
                'private String datasourceTag;                 // 데이터소스 태그',
                '',
                '@Column(name = "source_datasource_id")',
                'private String sourceDatasourceId;            // 소스 데이터소스 ID',
                '',
                '@Column(name = "target_datasource_id")',
                'private String targetDatasourceId;            // 타겟 데이터소스 ID',
                '',
                '@Column(name = "description", columnDefinition = "TEXT")',
                'private String description;                   // 설명',
                '',
                '@Enumerated(EnumType.STRING)',
                '@Column(name = "status")',
                'private AgentStatus status;                   // 상태 (ONLINE/OFFLINE/ERROR)',
                '',
                '@Column(name = "last_executed_at")',
                'private LocalDateTime lastExecutedAt;         // 마지막 실행 시각',
                '',
                '@Column(name = "last_execution_status")',
                'private String lastExecutionStatus;           // 마지막 실행 결과',
                '',
                '@CreationTimestamp',
                '@Column(name = "created_at", updatable = false)',
                'private LocalDateTime createdAt;              // 생성 시각',
                '',
                '@OneToMany(mappedBy = "agent", cascade = CascadeType.ALL, orphanRemoval = true)',
                'private List<Schedule> schedules;             // 스케줄 목록',
                '',
                '@OneToMany(mappedBy = "agent", cascade = CascadeType.ALL, orphanRemoval = true)',
                'private List<AgentChainMember> chainMembers;  // 체인 멤버 목록',
                '',
                '@OneToMany(mappedBy = "agent", cascade = CascadeType.ALL, orphanRemoval = true)',
                'private List<AgentTable> agentTables;         // Agent 테이블 목록',
                '',
                '@Column(name = "retention_config", columnDefinition = "TEXT")',
                'private String retentionConfig;               // Retention 설정 JSON',
            ]),
            'operations': None,
        },
        {
            'class_name': 'AgentTable',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'AgentTable 엔티티를 정의한다.',
                '테이블: agent_table. 에이전트-테이블 매핑 정보를 관리하며, Agent와 DatasourceTable 간의 SOURCE/TARGET 관계를 정의한다.'
            ),
            'attributes': '\n'.join([
                '@Id @GeneratedValue(strategy = GenerationType.IDENTITY)',
                'private Long id;                          // PK',
                '',
                '@ManyToOne(fetch = FetchType.LAZY)',
                '@JoinColumn(name = "agent_id", nullable = false)',
                'private Agent agent;                      // 에이전트 (FK)',
                '',
                '@Column(name = "datasource_table_id", nullable = false)',
                'private Long datasourceTableId;           // 데이터소스 테이블 FK',
                '',
                '@Enumerated(EnumType.STRING)',
                '@Column(name = "table_type", nullable = false)',
                'private TableType tableType;              // 테이블 유형 (SOURCE/TARGET)',
            ]),
            'operations': None,
        },
    ]))

    # --- AgentRepository ---
    agent_classes.append(('Repository', [{
        'class_name': 'AgentRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'AgentRepository를 정의한다.',
            'Agent 엔티티에 대한 JPA Repository로, 코드 조회, 존/상태별 필터링, 활성 Agent 조회, 상태별 카운트 기능을 제공한다.'
        ),
        'attributes': 'extends JpaRepository<Agent, Long>',
        'operations': '\n'.join([
            '// 에이전트 코드로 조회',
            'Optional<Agent> findByAgentCode(String agentCode);',
            '',
            '// 에이전트 코드 접두사로 조회',
            'List<Agent> findByAgentCodeStartingWith(String prefix);',
            '',
            '// 존별 조회',
            'List<Agent> findByZone(String zone);',
            '',
            '// 상태별 조회',
            'List<Agent> findByStatus(AgentStatus status);',
            '',
            '// 특정 상태 제외 조회',
            'List<Agent> findByStatusNot(AgentStatus status);',
            '',
            '// 존 + 특정 상태 제외 조회',
            'List<Agent> findByZoneAndStatusNot(String zone, AgentStatus status);',
            '',
            '// 활성화된 Agent만 조회',
            'List<Agent> findByIsActiveTrue();',
            '',
            '// 상태별 Agent 수 카운트',
            'long countByStatus(AgentStatus status);',
        ]),
    }]))

    groups.append(('Agent', agent_classes))

    # ================================================================
    # GROUP 2: Callback
    # ================================================================
    callback_classes = []

    callback_classes.append(('Controller', [{
        'class_name': 'CallbackController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'Callback Controller를 정의한다.',
            'Agent 실행 시작/완료 콜백 처리 기능을 제공한다.'
        ),
        'attributes': 'private final CallbackService callbackService;',
        'operations': '\n'.join([
            '// 실행 시작 콜백 처리',
            '@PostMapping("/started")',
            'public ResponseEntity<?> handleStarted(@RequestBody CallbackDto.StartedRequest request);',
            '',
            '// 실행 완료 콜백 처리',
            '@PostMapping("/finished")',
            'public ResponseEntity<?> handleFinished(@RequestBody CallbackDto.FinishedRequest request);',
        ]),
    }]))

    callback_classes.append(('Service', [{
        'class_name': 'CallbackService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'Callback 서비스를 정의한다.',
            'Agent 실행 시작/완료 콜백 처리, ExecutionHistory/StepHistory 생성/업데이트 비즈니스 로직을 처리한다.'
        ),
        'attributes': '\n'.join([
            'private final AgentRepository agentRepository;',
            'private final ExecutionHistoryRepository executionHistoryRepository;',
            'private final ExecutionStepHistoryRepository executionStepHistoryRepository;',
        ]),
        'operations': '\n'.join([
            '// 실행 시작 콜백 처리',
            'public void handleStarted(CallbackDto.StartedRequest request);',
            '',
            '// 실행 완료 콜백 처리',
            'public void handleFinished(CallbackDto.FinishedRequest request);',
        ]),
    }]))

    callback_classes.append(('Request', [{
        'class_name': 'CallbackRequest',
        'class_type': 'Class <<Request>>',
        'overview': make_overview(
            'Callback Request 정보를 정의한다.',
            '실행 시작/완료 콜백 요청에 필요한 실행ID, Agent ID, 상태, 통계, Step 결과 등의 파라미터를 정의한다.'
        ),
        'attributes': '\n'.join([
            'private String executionId;          // 실행 ID',
            'private String agentId;              // Agent ID',
            'private LocalDateTime startedAt;     // 실행 시작 시각',
            'private String triggeredBy;          // 실행 주체 (MANUAL/SCHEDULE/CHAIN)',
            'private String status;               // 실행 상태',
            'private Integer totalReadCount;      // 읽기 건수',
            'private Integer totalWriteCount;     // 쓰기 건수',
            'private Integer totalSkipCount;      // 스킵 건수',
            'private Long durationMs;             // 소요 시간 (ms)',
            'private String errorMessage;         // 오류 메시지',
            'private LocalDateTime finishedAt;    // 실행 종료 시각',
            'private List<StepResultItem> stepResults;  // Step 결과 목록',
            'private String stepId;               // 스텝 ID',
            'private Integer readCount;           // 스텝 읽기 건수',
            'private Integer writeCount;          // 스텝 쓰기 건수',
            'private Integer skipCount;           // 스텝 스킵 건수',
            'private Integer stepOrder;           // 스텝 순서',
        ]),
        'operations': None,
    }]))

    callback_classes.append(('Response', [{
        'class_name': 'CallbackResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'Callback Response 정보를 정의한다.',
            '콜백 처리 결과 응답은 공통 응답(HTTP 200) 처리되므로 별도 비즈니스 필드는 없다.'
        ),
        'attributes': '(없음)',
        'operations': None,
    }]))

    groups.append(('Callback', callback_classes))

    # ================================================================
    # GROUP 3: Chain (AgentChain)
    # ================================================================
    chain_classes = []

    chain_classes.append(('Controller', [{
        'class_name': 'AgentChainController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'AgentChain Controller를 정의한다.',
            'Agent 체인 CRUD 기능을 제공한다.'
        ),
        'attributes': 'private final AgentChainService chainService;',
        'operations': '\n'.join([
            '// 체인 전체 목록 조회',
            '@GetMapping',
            'public ResponseEntity<List<AgentChainDto.Response>> getChains();',
            '',
            '// 체인 단건 조회',
            '@GetMapping("/{chainId}")',
            'public ResponseEntity<AgentChainDto.Response> getChain(@PathVariable String chainId);',
            '',
            '// 체인 생성',
            '@PostMapping',
            'public ResponseEntity<AgentChainDto.Response> createChain(@Valid @RequestBody AgentChainDto.CreateRequest request);',
            '',
            '// 체인 수정',
            '@PutMapping("/{chainId}")',
            'public ResponseEntity<AgentChainDto.Response> updateChain(@PathVariable String chainId, @RequestBody AgentChainDto.UpdateRequest request);',
            '',
            '// 체인 삭제',
            '@DeleteMapping("/{chainId}")',
            'public ResponseEntity<Void> deleteChain(@PathVariable String chainId);',
        ]),
    }]))

    chain_classes.append(('Service', [{
        'class_name': 'AgentChainService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'AgentChain 서비스를 정의한다.',
            'Agent 체인 CRUD, 멤버 관리 비즈니스 로직을 처리한다.'
        ),
        'attributes': '\n'.join([
            'private final AgentChainRepository chainRepository;',
            'private final AgentRepository agentRepository;',
        ]),
        'operations': '\n'.join([
            '// 체인 전체 목록 조회',
            'public List<AgentChainDto.Response> findAll();',
            '',
            '// 체인 단건 조회',
            'public AgentChainDto.Response findById(String chainId);',
            '',
            '// 체인 생성',
            'public AgentChainDto.Response create(AgentChainDto.CreateRequest request);',
            '',
            '// 체인 수정',
            'public AgentChainDto.Response update(String chainId, AgentChainDto.UpdateRequest request);',
            '',
            '// 체인 삭제',
            'public void delete(String chainId);',
        ]),
    }]))

    chain_classes.append(('Request', [{
        'class_name': 'AgentChainRequest',
        'class_type': 'Class <<Request>>',
        'overview': make_overview(
            'AgentChain Request 정보를 정의한다.',
            '체인 생성/수정 요청에 필요한 체인ID, 이름, 설명, 트리거유형, 멤버 목록 등의 파라미터를 정의한다.'
        ),
        'attributes': '\n'.join([
            'private String chainId;              // 체인 ID',
            'private String chainName;            // 체인명',
            'private String description;          // 설명',
            'private TriggerType triggerType;     // 트리거 유형 (INDIVIDUAL/SEQUENTIAL)',
            'private List<MemberRequest> members; // 멤버 목록',
            'private Long agentId;                // Agent ID (멤버)',
            'private Integer seqOrder;            // 실행 순서 (멤버)',
        ]),
        'operations': None,
    }]))

    chain_classes.append(('Response', [{
        'class_name': 'AgentChainResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'AgentChain Response 정보를 정의한다.',
            '체인 조회 응답에 포함되는 체인ID, 이름, 트리거유형, 멤버 목록 등을 정의한다.'
        ),
        'attributes': '\n'.join([
            'private String chainId;                    // 체인 ID',
            'private String chainName;                  // 체인명',
            'private String description;                // 설명',
            'private TriggerType triggerType;           // 트리거 유형',
            'private List<MemberResponse> members;      // 멤버 목록',
            'private LocalDateTime createdAt;           // 생성 시각',
            'private Long id;                           // 멤버 ID',
            'private Long agentId;                      // Agent ID (멤버)',
            'private String agentCode;                  // Agent 코드 (멤버)',
            'private String agentName;                  // Agent 명 (멤버)',
            'private String zone;                       // 존 (멤버)',
            'private Integer seqOrder;                  // 실행 순서 (멤버)',
        ]),
        'operations': None,
    }]))

    chain_classes.append(('Entity', [
        {
            'class_name': 'AgentChain',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'AgentChain 엔티티를 정의한다.',
                '테이블: agent_chain. 에이전트 체인 정보를 관리하며, 체인ID, 이름, 트리거유형, 멤버 목록을 포함한다.'
            ),
            'attributes': '\n'.join([
                '@Id',
                '@Column(name = "chain_id")',
                'private String chainId;                   // 체인 ID (PK)',
                '',
                '@Column(name = "chain_name", nullable = false)',
                'private String chainName;                 // 체인명',
                '',
                '@Column(name = "description", columnDefinition = "TEXT")',
                'private String description;               // 설명',
                '',
                '@Enumerated(EnumType.STRING)',
                '@Column(name = "trigger_type", nullable = false)',
                'private TriggerType triggerType;          // 트리거 유형',
                '',
                '@OneToMany(mappedBy = "chain", cascade = CascadeType.ALL, orphanRemoval = true)',
                '@OrderBy("seqOrder ASC")',
                'private List<AgentChainMember> members;   // 체인 멤버 목록',
                '',
                '@CreationTimestamp',
                '@Column(name = "created_at", updatable = false)',
                'private LocalDateTime createdAt;          // 생성 시각',
            ]),
            'operations': None,
        },
        {
            'class_name': 'AgentChainMember',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'AgentChainMember 엔티티를 정의한다.',
                '테이블: agent_chain_member. 체인 구성원 정보를 관리하며, 체인-Agent 간 매핑 및 실행 순서를 정의한다.'
            ),
            'attributes': '\n'.join([
                '@Id @GeneratedValue(strategy = GenerationType.IDENTITY)',
                'private Long id;                          // PK',
                '',
                '@ManyToOne(fetch = FetchType.LAZY)',
                '@JoinColumn(name = "chain_id", nullable = false)',
                'private AgentChain chain;                 // 체인 (FK)',
                '',
                '@ManyToOne(fetch = FetchType.LAZY)',
                '@JoinColumn(name = "agent_id", nullable = false)',
                'private Agent agent;                      // 에이전트 (FK)',
                '',
                '@Column(name = "seq_order", nullable = false)',
                'private Integer seqOrder;                 // 실행 순서',
            ]),
            'operations': None,
        },
    ]))

    chain_classes.append(('Repository', [{
        'class_name': 'AgentChainRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'AgentChainRepository를 정의한다.',
            'AgentChain 엔티티에 대한 JPA Repository로, 멤버 포함 전체/단건 조회 기능을 제공한다.'
        ),
        'attributes': 'extends JpaRepository<AgentChain, String>',
        'operations': '\n'.join([
            '// 멤버 포함 전체 조회',
            '@Query("SELECT DISTINCT c FROM AgentChain c LEFT JOIN FETCH c.members m LEFT JOIN FETCH m.agent")',
            'List<AgentChain> findAllWithMembers();',
            '',
            '// 멤버 포함 단건 조회',
            '@Query("SELECT c FROM AgentChain c LEFT JOIN FETCH c.members m LEFT JOIN FETCH m.agent WHERE c.chainId = :chainId")',
            'Optional<AgentChain> findByIdWithMembers(@Param("chainId") String chainId);',
        ]),
    }]))

    groups.append(('Chain', chain_classes))

    # ================================================================
    # GROUP 4: Datasource
    # ================================================================
    ds_classes = []

    ds_classes.append(('Controller', [{
        'class_name': 'DatasourceController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'Datasource Controller를 정의한다.',
            'Datasource CRUD, 연결 테스트, 테이블/컬럼 검색, 테이블 등록/삭제, alias 맵 등의 기능을 제공한다.'
        ),
        'attributes': 'private final DatasourceService datasourceService;',
        'operations': '\n'.join([
            '// 전체 목록 조회',
            '@GetMapping',
            'public ResponseEntity<List<DatasourceDto.Response>> getDatasources();',
            '',
            '// 활성화된 목록 조회',
            '@GetMapping("/active")',
            'public ResponseEntity<List<DatasourceDto.Response>> getActiveDatasources();',
            '',
            '// 간단 목록 조회 (Agent 등록용)',
            '@GetMapping("/simple")',
            'public ResponseEntity<List<DatasourceDto.SimpleResponse>> getDatasourcesSimple();',
            '',
            '// sourceRef 해석용 lookup 조회',
            '@GetMapping("/sourceref-lookup")',
            'public ResponseEntity<DatasourceDto.SourceRefLookup> getSourceRefLookup();',
            '',
            '// 테이블 alias 전역 조회',
            '@GetMapping("/table-alias-map")',
            'public ResponseEntity<Map<String, String>> getTableAliasMap();',
            '',
            '// 단건 조회',
            '@GetMapping("/{datasourceId}")',
            'public ResponseEntity<DatasourceDto.Response> getDatasource(@PathVariable String datasourceId);',
            '',
            '// 연결 정보 조회 (Agent 내부용)',
            '@GetMapping("/{datasourceId}/connection-info")',
            'public ResponseEntity<DatasourceDto.ConnectionInfo> getConnectionInfo(@PathVariable String datasourceId);',
            '',
            '// 생성',
            '@PostMapping',
            'public ResponseEntity<DatasourceDto.Response> createDatasource(@Valid @RequestBody DatasourceDto.CreateRequest request);',
            '',
            '// 수정',
            '@PutMapping("/{datasourceId}")',
            'public ResponseEntity<DatasourceDto.Response> updateDatasource(@PathVariable String datasourceId, @RequestBody DatasourceDto.UpdateRequest request);',
            '',
            '// 삭제',
            '@DeleteMapping("/{datasourceId}")',
            'public ResponseEntity<Void> deleteDatasource(@PathVariable String datasourceId);',
            '',
            '// 저장된 datasource 연결 테스트',
            '@PostMapping("/{datasourceId}/test-connection")',
            'public ResponseEntity<DatasourceDto.ConnectionTestResponse> testConnection(@PathVariable String datasourceId);',
            '',
            '// 저장 전 연결 테스트',
            '@PostMapping("/test-connection")',
            'public ResponseEntity<DatasourceDto.ConnectionTestResponse> testConnectionBeforeSave(@Valid @RequestBody DatasourceDto.ConnectionTestRequest request);',
            '',
            '// 테이블 검색',
            '@GetMapping("/{datasourceId}/search-tables")',
            'public ResponseEntity<List<DatasourceDto.TableSearchResult>> searchTables(@PathVariable String datasourceId, @RequestParam(required=false) String query);',
            '',
            '// 컬럼 검색',
            '@GetMapping("/{datasourceId}/search-columns")',
            'public ResponseEntity<List<DatasourceDto.ColumnSearchResult>> searchColumns(@PathVariable String datasourceId, @RequestParam String tableName, @RequestParam(required=false) String query);',
            '',
            '// 등록된 테이블 목록 조회',
            '@GetMapping("/{datasourceId}/tables")',
            'public ResponseEntity<List<DatasourceDto.TableResponse>> getRegisteredTables(@PathVariable String datasourceId);',
            '',
            '// 테이블 등록',
            '@PostMapping("/{datasourceId}/tables")',
            'public ResponseEntity<DatasourceDto.TableResponse> registerTable(@PathVariable String datasourceId, @Valid @RequestBody DatasourceDto.TableCreateRequest request);',
            '',
            '// 테이블 삭제',
            '@DeleteMapping("/{datasourceId}/tables/{tableId}")',
            'public ResponseEntity<Void> deleteTable(@PathVariable String datasourceId, @PathVariable Long tableId);',
        ]),
    }]))

    ds_classes.append(('Service', [{
        'class_name': 'DatasourceService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'Datasource 서비스를 정의한다.',
            'Datasource CRUD, 암호화 처리, 연결 테스트(zone별 프록시 분기), 테이블/컬럼 검색, 테이블 등록/삭제 비즈니스 로직을 처리한다.'
        ),
        'attributes': '\n'.join([
            'private final DatasourceRepository datasourceRepository;',
            'private final DatasourceTableRepository tableRepository;',
            'private final AgentRepository agentRepository;',
            'private final ZoneConfigRepository zoneConfigRepository;',
            'private final RestTemplate restTemplate;',
            'private final PasswordEncryptor passwordEncryptor;',
        ]),
        'operations': '\n'.join([
            '// 전체 목록 조회',
            'public List<DatasourceDto.Response> findAll();',
            '',
            '// 활성화된 목록 조회',
            'public List<DatasourceDto.Response> findActive();',
            '',
            '// 연결 정보 조회 (암호문 전달)',
            'public DatasourceDto.ConnectionInfo getConnectionInfo(String datasourceId);',
            '',
            '// 간단 목록 조회',
            'public List<DatasourceDto.SimpleResponse> findAllSimple();',
            '',
            '// 테이블 alias 전역 조회',
            'public Map<String, String> getTableAliasMap();',
            '',
            '// sourceRef 해석용 lookup 조회',
            'public DatasourceDto.SourceRefLookup getSourceRefLookup();',
            '',
            '// 단건 조회',
            'public DatasourceDto.Response findById(String datasourceId);',
            '',
            '// 생성 (암호화 처리)',
            'public DatasourceDto.Response create(DatasourceDto.CreateRequest request);',
            '',
            '// 수정',
            'public DatasourceDto.Response update(String datasourceId, DatasourceDto.UpdateRequest request);',
            '',
            '// 삭제',
            'public void delete(String datasourceId);',
            '',
            '// 저장된 datasource 연결 테스트',
            'public DatasourceDto.ConnectionTestResponse testConnection(String datasourceId);',
            '',
            '// 입력값 직접 연결 테스트',
            'public DatasourceDto.ConnectionTestResponse testConnection(DatasourceDto.ConnectionTestRequest request);',
            '',
            '// 테이블 검색',
            'public List<DatasourceDto.TableSearchResult> searchTables(String datasourceId, String query);',
            '',
            '// 컬럼 검색',
            'public List<DatasourceDto.ColumnSearchResult> searchColumns(String datasourceId, String tableName, String query);',
            '',
            '// 등록된 테이블 목록 조회',
            'public List<DatasourceDto.TableResponse> getRegisteredTables(String datasourceId);',
            '',
            '// 테이블 등록',
            'public DatasourceDto.TableResponse registerTable(String datasourceId, DatasourceDto.TableCreateRequest request);',
            '',
            '// 테이블 삭제',
            'public void deleteTable(String datasourceId, Long tableId);',
        ]),
    }]))

    ds_classes.append(('Request', [{
        'class_name': 'DatasourceRequest',
        'class_type': 'Class <<Request>>',
        'overview': make_overview(
            'Datasource Request 정보를 정의한다.',
            'Datasource 생성/수정, 연결 테스트, 테이블/컬럼 등록 요청에 필요한 파라미터를 정의한다.'
        ),
        'attributes': '\n'.join([
            'private String datasourceId;        // 데이터소스 ID',
            'private String datasourceName;      // 데이터소스명',
            'private DbType dbType;              // DB 유형',
            'private String host;                // 호스트 주소',
            'private Integer port;               // 포트 번호',
            'private String databaseName;        // 데이터베이스명',
            'private String username;            // 접속 계정',
            'private String password;            // 접속 비밀번호',
            'private String description;         // 설명',
            'private String zone;                // 네트워크 존',
            'private Boolean isActive;           // 활성화 여부',
            'private String tableName;           // 테이블명',
            'private String tableAlias;          // 테이블 별칭',
            'private List<ColumnCreateRequest> columns;  // 컬럼 목록',
            'private String columnName;          // 컬럼명',
            'private String columnAlias;         // 컬럼 별칭',
            'private String dataType;            // 데이터 타입',
            'private Boolean isPrimaryKey;       // PK 여부',
            'private Boolean isNullable;         // NULL 허용 여부',
        ]),
        'operations': None,
    }]))

    ds_classes.append(('Response', [{
        'class_name': 'DatasourceResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'Datasource Response 정보를 정의한다.',
            'Datasource 조회, 연결 테스트, 테이블/컬럼 검색, sourceRef lookup 등의 응답 데이터를 정의한다.'
        ),
        'attributes': '\n'.join([
            'private Long id;                         // 자동생성 ID',
            'private String datasourceId;             // 데이터소스 ID',
            'private String datasourceName;           // 데이터소스명',
            'private DbType dbType;                   // DB 유형',
            'private String host;                     // 호스트 주소',
            'private Integer port;                    // 포트 번호',
            'private String databaseName;             // 데이터베이스명',
            'private String username;                 // 접속 계정',
            'private String password;                 // 접속 비밀번호 (ConnectionInfo 용)',
            'private String description;              // 설명',
            'private String zone;                     // 네트워크 존',
            'private Boolean isActive;                // 활성화 여부',
            'private LocalDateTime createdAt;         // 생성 시각',
            'private LocalDateTime updatedAt;         // 수정 시각',
            'private boolean success;                 // 연결 테스트 결과',
            'private String message;                  // 연결 테스트 메시지',
            'private Long responseTimeMs;             // 응답 시간 (ms)',
            'private Map<Long, String> datasources;   // sourceRef lookup - 데이터소스 맵',
            'private Map<Long, String> tables;        // sourceRef lookup - 테이블 맵',
            'private String tableName;                // 테이블명',
            'private String tableType;                // 테이블 유형 (TABLE/VIEW)',
            'private String remarks;                  // 비고',
            'private String columnName;               // 컬럼명',
            'private String dataType;                 // 데이터 타입',
            'private Integer columnSize;              // 컬럼 크기',
            'private Boolean isNullable;              // NULL 허용 여부',
            'private Boolean isPrimaryKey;            // PK 여부',
            'private String tableAlias;               // 테이블 별칭',
            'private List<ColumnResponse> columns;    // 컬럼 목록',
            'private String columnAlias;              // 컬럼 별칭',
        ]),
        'operations': None,
    }]))

    ds_classes.append(('Entity', [
        {
            'class_name': 'Datasource',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'Datasource 엔티티를 정의한다.',
                '테이블: datasource. 데이터소스 연결 정보를 관리하며, DB유형/호스트/포트/계정(암호화)/존 등을 포함한다.'
            ),
            'attributes': '\n'.join([
                '@GeneratedValue(strategy = GenerationType.IDENTITY)',
                '@Column(name = "id", insertable = false, updatable = false)',
                'private Long id;                          // 자동생성 시퀀스',
                '',
                '@Id',
                '@Column(name = "datasource_id")',
                'private String datasourceId;              // 데이터소스 고유 ID (PK)',
                '',
                '@Column(name = "datasource_name", nullable = false)',
                'private String datasourceName;            // 데이터소스명',
                '',
                '@Enumerated(EnumType.STRING)',
                '@Column(name = "db_type", nullable = false)',
                'private DbType dbType;                    // DB 유형',
                '',
                '@Column(name = "host", nullable = false)',
                'private String host;                      // 호스트 주소',
                '',
                '@Column(name = "port", nullable = false)',
                'private Integer port;                     // 포트 번호',
                '',
                '@Column(name = "database_name", nullable = false)',
                'private String databaseName;              // 데이터베이스명',
                '',
                '@Column(name = "username", nullable = false)',
                'private String username;                  // 접속 계정 (암호화)',
                '',
                '@Column(name = "password", nullable = false)',
                'private String password;                  // 접속 비밀번호 (암호화)',
                '',
                '@Column(name = "description", columnDefinition = "TEXT")',
                'private String description;               // 설명',
                '',
                '@Column(name = "zone")',
                'private String zone;                      // 네트워크 존',
                '',
                '@Column(name = "is_active")',
                'private Boolean isActive;                 // 활성화 여부',
                '',
                '@CreationTimestamp',
                '@Column(name = "created_at", updatable = false)',
                'private LocalDateTime createdAt;          // 생성 시각',
                '',
                '@UpdateTimestamp',
                '@Column(name = "updated_at")',
                'private LocalDateTime updatedAt;          // 수정 시각',
            ]),
            'operations': None,
        },
        {
            'class_name': 'DatasourceTable',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'DatasourceTable 엔티티를 정의한다.',
                '테이블: datasource_table. 데이터소스에 등록된 테이블 정보를 관리하며, 테이블명/별칭/컬럼 목록을 포함한다.'
            ),
            'attributes': '\n'.join([
                '@Id @GeneratedValue(strategy = GenerationType.IDENTITY)',
                'private Long id;                          // PK',
                '',
                '@Column(name = "datasource_id", nullable = false)',
                'private String datasourceId;              // 데이터소스 ID',
                '',
                '@Column(name = "table_name", nullable = false)',
                'private String tableName;                 // 테이블명',
                '',
                '@Column(name = "table_alias")',
                'private String tableAlias;                // 테이블 별칭',
                '',
                '@Column(name = "description")',
                'private String description;               // 설명',
                '',
                '@OneToMany(mappedBy = "datasourceTable", cascade = CascadeType.ALL, orphanRemoval = true)',
                'private List<DatasourceColumn> columns;   // 컬럼 목록',
                '',
                '@CreationTimestamp',
                '@Column(name = "created_at", updatable = false)',
                'private LocalDateTime createdAt;          // 생성 시각',
            ]),
            'operations': None,
        },
        {
            'class_name': 'DatasourceColumn',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'DatasourceColumn 엔티티를 정의한다.',
                '테이블: datasource_column. 데이터소스 테이블의 컬럼 정보를 관리하며, 컬럼명/별칭/데이터타입/PK여부 등을 포함한다.'
            ),
            'attributes': '\n'.join([
                '@Id @GeneratedValue(strategy = GenerationType.IDENTITY)',
                'private Long id;                          // PK',
                '',
                '@ManyToOne(fetch = FetchType.LAZY)',
                '@JoinColumn(name = "datasource_table_id", nullable = false)',
                'private DatasourceTable datasourceTable;  // 데이터소스 테이블 (FK)',
                '',
                '@Column(name = "column_name", nullable = false)',
                'private String columnName;                // 컬럼명',
                '',
                '@Column(name = "column_alias")',
                'private String columnAlias;               // 컬럼 별칭',
                '',
                '@Column(name = "data_type")',
                'private String dataType;                  // 데이터 타입',
                '',
                '@Column(name = "is_primary_key")',
                'private Boolean isPrimaryKey;             // PK 여부',
                '',
                '@Column(name = "is_nullable")',
                'private Boolean isNullable;               // NULL 허용 여부',
                '',
                '@Column(name = "description")',
                'private String description;               // 설명',
            ]),
            'operations': None,
        },
    ]))

    ds_classes.append(('Repository', [
        {
            'class_name': 'DatasourceRepository',
            'class_type': 'Interface <<Repository>>',
            'overview': make_overview(
                'DatasourceRepository를 정의한다.',
                'Datasource 엔티티에 대한 JPA Repository로, 활성 조회, DB유형별 조회 기능을 제공한다.'
            ),
            'attributes': 'extends JpaRepository<Datasource, String>',
            'operations': '\n'.join([
                '// 활성화된 목록 조회',
                'List<Datasource> findByIsActiveTrue();',
                '',
                '// DB 유형별 조회',
                'List<Datasource> findByDbType(DbType dbType);',
                '',
                '// DB 유형 + 활성화 조회',
                'List<Datasource> findByDbTypeAndIsActiveTrue(DbType dbType);',
            ]),
        },
        {
            'class_name': 'DatasourceTableRepository',
            'class_type': 'Interface <<Repository>>',
            'overview': make_overview(
                'DatasourceTableRepository를 정의한다.',
                'DatasourceTable 엔티티에 대한 JPA Repository로, 데이터소스별 테이블 조회, 중복 검사 기능을 제공한다.'
            ),
            'attributes': 'extends JpaRepository<DatasourceTable, Long>',
            'operations': '\n'.join([
                '// 데이터소스별 테이블 조회 (컬럼 포함)',
                '@EntityGraph(attributePaths = {"columns"})',
                'List<DatasourceTable> findByDatasourceId(String datasourceId);',
                '',
                '// 데이터소스 + 테이블명으로 조회',
                'Optional<DatasourceTable> findByDatasourceIdAndTableName(String datasourceId, String tableName);',
                '',
                '// 데이터소스 + 테이블명 존재 여부',
                'boolean existsByDatasourceIdAndTableName(String datasourceId, String tableName);',
            ]),
        },
    ]))

    groups.append(('Datasource', ds_classes))

    # ================================================================
    # GROUP 5: Execution
    # ================================================================
    exec_classes = []

    exec_classes.append(('Controller', [{
        'class_name': 'ExecutionController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'Execution Controller를 정의한다.',
            '실행 트리거, 실행 상세/데이터 조회, 테이블별 통계, 데이터 추적, Step 로그, 이력 조회, 대시보드 통계 기능을 제공한다.'
        ),
        'attributes': 'private final ExecutionService executionService;',
        'operations': '\n'.join([
            '// Agent별 실행 이력 조회 (Agent DB)',
            '@GetMapping("/agent/{id}")',
            'public ResponseEntity<List<Map<String, Object>>> getExecutionsByAgent(@PathVariable Long id);',
            '',
            '// 전체 Agent 상태 조회 (대시보드용)',
            '@GetMapping("/status")',
            'public ResponseEntity<List<ExecutionDto.AgentExecutionSummary>> getAgentStatuses();',
            '',
            '// 실행 트리거',
            '@PostMapping("/{id}/run")',
            'public ResponseEntity<ExecutionDto.TriggerResponse> triggerExecution(@PathVariable Long id, @RequestBody(required=false) ExecutionDto.TriggerRequest request);',
            '',
            '// 실행 상세 정보 조회',
            '@GetMapping("/{executionId}/detail")',
            'public ResponseEntity<Map<String, Object>> getExecutionDetail(@PathVariable String executionId);',
            '',
            '// 실행 데이터 조회 (페이징/검색)',
            '@GetMapping("/{executionId}/data/{dataType}")',
            'public ResponseEntity<Map<String, Object>> getExecutionData(@PathVariable String executionId, @PathVariable String dataType, ExecutionDto.TableDataSearchParams searchParams);',
            '',
            '// 테이블별 통계 조회',
            '@GetMapping("/{executionId}/tables")',
            'public ResponseEntity<List<Map<String, Object>>> getTableStats(@PathVariable String executionId);',
            '',
            '// 특정 테이블 레코드 조회',
            '@GetMapping("/{executionId}/tables/{tableName}")',
            'public ResponseEntity<Map<String, Object>> getTableRecords(@PathVariable String executionId, @PathVariable String tableName);',
            '',
            '// 특정 테이블 실패 레코드 조회',
            '@GetMapping("/{executionId}/tables/{tableName}/failed")',
            'public ResponseEntity<List<Map<String, Object>>> getTableFailedRecords(@PathVariable String executionId, @PathVariable String tableName);',
            '',
            '// Source PK로 데이터 추적',
            '@GetMapping("/{executionId}/trace")',
            'public ResponseEntity<Map<String, Object>> traceBySourcePk(@PathVariable String executionId, @RequestParam String pkValue, @RequestParam(defaultValue="id") String pkColumn, @RequestParam String sourceTable, @RequestParam(required=false) String ifTableName, @RequestParam(required=false) String targetTableName);',
            '',
            '// Target에서 Source로 역추적',
            '@GetMapping("/{executionId}/trace-source")',
            'public ResponseEntity<Map<String, Object>> traceToSource(@PathVariable String executionId, @RequestParam String sourceRefs, @RequestParam(required=false) String sourceTable);',
            '',
            '// Step별 결과 조회',
            '@GetMapping("/{executionId}/steps")',
            'public ResponseEntity<List<ExecutionStepHistory>> getExecutionSteps(@PathVariable String executionId);',
            '',
            '// 최근 실행 이력 조회',
            '@GetMapping("/history")',
            'public ResponseEntity<List<ExecutionDto.HistoryResponse>> getRecentHistory();',
            '',
            '// 실행 중인 이력 조회',
            '@GetMapping("/history/running")',
            'public ResponseEntity<List<ExecutionDto.HistoryResponse>> getRunningHistory();',
            '',
            '// Agent별 실행 이력 조회 (Orchestrator DB)',
            '@GetMapping("/history/agent/{id}")',
            'public ResponseEntity<List<ExecutionDto.HistoryResponse>> getHistoryByAgent(@PathVariable Long id);',
            '',
            '// 실행 이력 페이징 조회',
            '@GetMapping("/history/paged")',
            'public ResponseEntity<Page<ExecutionDto.HistoryResponse>> getHistoryPaged(@RequestParam(defaultValue="0") int page, @RequestParam(defaultValue="20") int size, @RequestParam(required=false) String status, @RequestParam(required=false) String agentCode, @RequestParam(required=false) String agentType, @RequestParam(required=false) String zone, @RequestParam(required=false) String startDate, @RequestParam(required=false) String endDate, @RequestParam(required=false) String search);',
            '',
            '// 대시보드 통계 조회',
            '@GetMapping("/dashboard/stats")',
            'public ResponseEntity<ExecutionDto.DashboardStats> getDashboardStats();',
        ]),
    }]))

    exec_classes.append(('Service', [{
        'class_name': 'ExecutionService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'Execution 서비스를 정의한다.',
            '실행 트리거(Agent 호출), 실행 상세/데이터/통계 조회(프록시 경유), 데이터 추적, Step 로그, 이력 조회/필터링, 대시보드 통계 비즈니스 로직을 처리한다.'
        ),
        'attributes': '\n'.join([
            'private final AgentRepository agentRepository;',
            'private final DatasourceRepository datasourceRepository;',
            'private final DatasourceTableRepository datasourceTableRepository;',
            'private final ZoneConfigRepository zoneConfigRepository;',
            'private final ExecutionHistoryRepository executionHistoryRepository;',
            'private final ExecutionStepHistoryRepository executionStepHistoryRepository;',
            'private final PasswordEncryptor passwordEncryptor;',
            'private final RestTemplate restTemplate;',
        ]),
        'operations': '\n'.join([
            '// Agent별 실행 이력 조회 (프록시 경유)',
            'public List<Map<String, Object>> findByAgentIdFromAgent(Long id);',
            '',
            '// 전체 Agent 상태 조회',
            'public List<ExecutionDto.AgentExecutionSummary> getAgentStatuses();',
            '',
            '// 실행 상세 정보 조회',
            'public Map<String, Object> getExecutionDetail(String executionId);',
            '',
            '// 실행 데이터 조회 (페이징)',
            'public Map<String, Object> getExecutionData(String executionId, String dataType, ExecutionDto.TableDataSearchParams searchParams);',
            '',
            '// 테이블별 통계 조회',
            'public List<Map<String, Object>> getTableStats(String executionId);',
            '',
            '// 특정 테이블 레코드 조회',
            'public Map<String, Object> getTableRecords(String executionId, String tableName);',
            '',
            '// 특정 테이블 실패 레코드 조회',
            'public List<Map<String, Object>> getTableFailedRecords(String executionId, String tableName);',
            '',
            '// Source PK로 데이터 추적',
            'public Map<String, Object> traceBySourcePk(String executionId, String pkValue, String pkColumn, String sourceTable, String ifTableName, String targetTableName);',
            '',
            '// Target에서 Source로 역추적',
            'public Map<String, Object> traceToSource(String executionId, String sourceRefs, String sourceTable);',
            '',
            '// 실행 트리거 (내부 구현)',
            'public ExecutionDto.TriggerResponse triggerExecutionInternal(Long id, LocalDateTime startTime, LocalDateTime endTime, List<Map<String,Object>> filters, List<String> selectedStepIds, List<Map<String,Object>> conditions, String triggeredBy);',
            '',
            '// Step별 결과 조회',
            'public List<ExecutionStepHistory> getExecutionSteps(String executionId);',
            '',
            '// 최근 실행 이력 조회',
            'public List<ExecutionDto.HistoryResponse> getRecentHistory();',
            '',
            '// 실행 중인 이력 조회',
            'public List<ExecutionDto.HistoryResponse> getRunningHistory();',
            '',
            '// Agent별 실행 이력 조회',
            'public List<ExecutionDto.HistoryResponse> getHistoryByAgent(Long id);',
            '',
            '// 실행 이력 페이징 조회 (필터/검색)',
            'public Page<ExecutionDto.HistoryResponse> getHistoryPaged(int page, int size, String status, String agentCode, String agentType, String zone, String startDate, String endDate, String search);',
            '',
            '// 대시보드 통계 조회',
            'public ExecutionDto.DashboardStats getDashboardStats();',
        ]),
    }]))

    exec_classes.append(('Request', [{
        'class_name': 'ExecutionRequest',
        'class_type': 'Class <<Request>>',
        'overview': make_overview(
            'Execution Request 정보를 정의한다.',
            '실행 트리거 요청에 필요한 시간 범위, 필터, Step 선택, 조건 등의 파라미터와 데이터 조회 검색 파라미터를 정의한다.'
        ),
        'attributes': '\n'.join([
            'private LocalDateTime startTime;               // 시작 시간',
            'private LocalDateTime endTime;                 // 종료 시간',
            'private List<Map<String, Object>> filters;     // 필터 목록',
            'private List<String> selectedStepIds;          // 선택된 Step ID 목록',
            'private List<Map<String, Object>> conditions;  // 실행 조건 목록',
            'private int page;                              // 페이지 번호',
            'private int size;                              // 페이지 크기',
            'private String search;                         // 검색어',
            'private String searchColumn;                   // 검색 컬럼',
            'private String status;                         // 상태 필터',
            'private String tableName;                      // 테이블명 필터',
            'private String sortColumn;                     // 정렬 컬럼',
            'private String sortDirection;                  // 정렬 방향',
        ]),
        'operations': None,
    }]))

    exec_classes.append(('Response', [{
        'class_name': 'ExecutionResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'Execution Response 정보를 정의한다.',
            '실행 트리거 결과, Agent 상태 요약, 실행 이력, 대시보드 통계 등의 응답 데이터를 정의한다.'
        ),
        'attributes': '\n'.join([
            'private String executionId;                    // 실행 ID',
            'private Long agentId;                          // Agent ID',
            'private String agentCode;                      // Agent 코드',
            'private String agentName;                      // Agent 명',
            'private String zone;                           // 네트워크 존',
            'private String status;                         // 실행 상태',
            'private LocalDateTime startTime;               // 시작 시간',
            'private LocalDateTime endTime;                 // 종료 시간',
            'private ExecutionStatus lastExecutionStatus;   // 마지막 실행 상태',
            'private LocalDateTime lastRunAt;               // 마지막 실행 시각',
            'private AgentStatus agentStatus;               // Agent 상태',
            'private String agentType;                      // Agent 유형',
            'private Long totalReadCount;                   // 읽기 건수',
            'private Long totalWriteCount;                  // 쓰기 건수',
            'private Long totalSkipCount;                   // 스킵 건수',
            'private Long durationMs;                       // 소요 시간 (ms)',
            'private String errorMessage;                   // 오류 메시지',
            'private LocalDateTime startedAt;               // 실행 시작 시각',
            'private LocalDateTime finishedAt;              // 실행 종료 시각',
            'private String triggeredBy;                    // 실행 주체',
            'private long todayExecutions;                  // 오늘 실행 수',
            'private long todayFailed;                      // 오늘 실패 수',
            'private long currentlyRunning;                 // 현재 실행 중',
            'private long totalAgents;                      // 전체 Agent 수',
            'private long onlineAgents;                     // 온라인 Agent 수',
        ]),
        'operations': None,
    }]))

    exec_classes.append(('Entity', [
        {
            'class_name': 'ExecutionHistory',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'ExecutionHistory 엔티티를 정의한다.',
                '테이블: execution_history. 실행 이력 정보를 관리하며, 실행ID/Agent코드/상태/통계/시간/오류 등을 포함한다. 인덱스: agent_code, status, started_at DESC.'
            ),
            'attributes': '\n'.join([
                '@Id',
                '@Column(name = "execution_id")',
                'private String executionId;               // 실행 ID (PK)',
                '',
                '@Column(name = "agent_code", nullable = false)',
                'private String agentCode;                 // 에이전트 코드',
                '',
                '@Column(name = "agent_name")',
                'private String agentName;                 // 에이전트명',
                '',
                '@Enumerated(EnumType.STRING)',
                '@Column(name = "status")',
                'private ExecutionStatus status;           // 실행 상태',
                '',
                '@Column(name = "total_read_count")',
                'private Long totalReadCount;              // 읽기 건수',
                '',
                '@Column(name = "total_write_count")',
                'private Long totalWriteCount;             // 쓰기 건수',
                '',
                '@Column(name = "total_skip_count")',
                'private Long totalSkipCount;              // 스킵 건수',
                '',
                '@Column(name = "duration_ms")',
                'private Long durationMs;                  // 소요 시간 (ms)',
                '',
                '@Column(name = "error_message", columnDefinition = "TEXT")',
                'private String errorMessage;              // 오류 메시지',
                '',
                '@Column(name = "started_at")',
                'private LocalDateTime startedAt;          // 실행 시작 시각',
                '',
                '@Column(name = "finished_at")',
                'private LocalDateTime finishedAt;         // 실행 종료 시각',
                '',
                '@Column(name = "triggered_by")',
                'private String triggeredBy;               // 실행 주체 (MANUAL/SCHEDULE/CHAIN)',
                '',
                '@Column(name = "agent_type")',
                'private String agentType;                 // 에이전트 유형',
            ]),
            'operations': None,
        },
        {
            'class_name': 'ExecutionStepHistory',
            'class_type': 'Class <<Entity>>',
            'overview': make_overview(
                'ExecutionStepHistory 엔티티를 정의한다.',
                '테이블: execution_step_history. 실행 스텝별 이력을 관리하며, 스텝ID/상태/통계/순서 등을 포함한다. 인덱스: execution_id.'
            ),
            'attributes': '\n'.join([
                '@Id @GeneratedValue(strategy = GenerationType.IDENTITY)',
                'private Long id;                          // PK',
                '',
                '@Column(name = "execution_id", nullable = false)',
                'private String executionId;               // 실행 이력 FK',
                '',
                '@Column(name = "step_id", nullable = false)',
                'private String stepId;                    // 스텝 ID',
                '',
                '@Column(name = "status")',
                'private String status;                    // 스텝 상태',
                '',
                '@Column(name = "read_count")',
                'private Integer readCount;                // 읽기 건수',
                '',
                '@Column(name = "write_count")',
                'private Integer writeCount;               // 쓰기 건수',
                '',
                '@Column(name = "skip_count")',
                'private Integer skipCount;                // 스킵 건수',
                '',
                '@Column(name = "duration_ms")',
                'private Long durationMs;                  // 소요 시간 (ms)',
                '',
                '@Column(name = "error_message", columnDefinition = "TEXT")',
                'private String errorMessage;              // 오류 메시지',
                '',
                '@Column(name = "step_order")',
                'private Integer stepOrder;                // 스텝 순서',
            ]),
            'operations': None,
        },
    ]))

    exec_classes.append(('Repository', [
        {
            'class_name': 'ExecutionHistoryRepository',
            'class_type': 'Interface <<Repository>>',
            'overview': make_overview(
                'ExecutionHistoryRepository를 정의한다.',
                'ExecutionHistory 엔티티에 대한 JPA Repository로, 최근 이력, Agent별 조회, 상태별 조회, 페이징, 오늘자 통계 기능을 제공한다.'
            ),
            'attributes': 'extends JpaRepository<ExecutionHistory, String>, JpaSpecificationExecutor<ExecutionHistory>',
            'operations': '\n'.join([
                '// 최근 50건 조회',
                'List<ExecutionHistory> findTop50ByOrderByStartedAtDesc();',
                '',
                '// Agent별 이력 조회',
                'List<ExecutionHistory> findByAgentCodeOrderByStartedAtDesc(String agentCode);',
                '',
                '// Agent별 최근 10건 조회',
                'List<ExecutionHistory> findTop10ByAgentCodeOrderByStartedAtDesc(String agentCode);',
                '',
                '// 상태별 조회 (내림차순)',
                'List<ExecutionHistory> findByStatusOrderByStartedAtDesc(ExecutionStatus status);',
                '',
                '// 기간별 조회',
                'List<ExecutionHistory> findByStartedAtBetweenOrderByStartedAtDesc(LocalDateTime start, LocalDateTime end);',
                '',
                '// 상태별 조회 (오름차순)',
                'List<ExecutionHistory> findByStatusOrderByStartedAtAsc(ExecutionStatus status);',
                '',
                '// 전체 페이징 조회',
                'Page<ExecutionHistory> findAllByOrderByStartedAtDesc(Pageable pageable);',
                '',
                '// Agent별 페이징 조회',
                'Page<ExecutionHistory> findByAgentCodeOrderByStartedAtDesc(String agentCode, Pageable pageable);',
                '',
                '// 오늘 실행 건수',
                '@Query("SELECT COUNT(e) FROM ExecutionHistory e WHERE e.startedAt >= :today")',
                'long countTodayExecutions(@Param("today") LocalDateTime today);',
                '',
                '// 오늘 실패 건수',
                '@Query("SELECT COUNT(e) FROM ExecutionHistory e WHERE e.startedAt >= :today AND e.status = \'FAILED\'")',
                'long countTodayFailedExecutions(@Param("today") LocalDateTime today);',
                '',
                '// 상태별 카운트',
                'long countByStatus(ExecutionStatus status);',
            ]),
        },
        {
            'class_name': 'ExecutionStepHistoryRepository',
            'class_type': 'Interface <<Repository>>',
            'overview': make_overview(
                'ExecutionStepHistoryRepository를 정의한다.',
                'ExecutionStepHistory 엔티티에 대한 JPA Repository로, 실행 ID별 스텝 결과 순서대로 조회하는 기능을 제공한다.'
            ),
            'attributes': 'extends JpaRepository<ExecutionStepHistory, Long>',
            'operations': '\n'.join([
                '// 실행 ID별 스텝 순서대로 조회',
                'List<ExecutionStepHistory> findByExecutionIdOrderByStepOrder(String executionId);',
            ]),
        },
    ]))

    groups.append(('Execution', exec_classes))

    # ================================================================
    # GROUP 6: Schedule
    # ================================================================
    sched_classes = []

    sched_classes.append(('Controller', [{
        'class_name': 'ScheduleController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'Schedule Controller를 정의한다.',
            '스케줄 CRUD, 활성화/비활성화 토글 기능을 제공한다.'
        ),
        'attributes': 'private final ScheduleService scheduleService;',
        'operations': '\n'.join([
            '// 스케줄 전체 목록 조회',
            '@GetMapping',
            'public ResponseEntity<List<ScheduleDto.Response>> getSchedules();',
            '',
            '// 스케줄 단건 조회',
            '@GetMapping("/{scheduleId}")',
            'public ResponseEntity<ScheduleDto.Response> getSchedule(@PathVariable Long scheduleId);',
            '',
            '// 스케줄 생성',
            '@PostMapping',
            'public ResponseEntity<ScheduleDto.Response> createSchedule(@Valid @RequestBody ScheduleDto.CreateRequest request);',
            '',
            '// 스케줄 수정',
            '@PutMapping("/{scheduleId}")',
            'public ResponseEntity<ScheduleDto.Response> updateSchedule(@PathVariable Long scheduleId, @RequestBody ScheduleDto.UpdateRequest request);',
            '',
            '// 스케줄 활성화/비활성화 토글',
            '@PutMapping("/{scheduleId}/toggle")',
            'public ResponseEntity<ScheduleDto.Response> toggleSchedule(@PathVariable Long scheduleId);',
            '',
            '// 스케줄 삭제',
            '@DeleteMapping("/{scheduleId}")',
            'public ResponseEntity<Void> deleteSchedule(@PathVariable Long scheduleId);',
        ]),
    }]))

    sched_classes.append(('Service', [{
        'class_name': 'ScheduleService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'Schedule 서비스를 정의한다.',
            '스케줄 CRUD, ScheduleExecutor 등록/갱신/취소, 활성화/비활성화 토글 비즈니스 로직을 처리한다.'
        ),
        'attributes': '\n'.join([
            'private final ScheduleRepository scheduleRepository;',
            'private final AgentRepository agentRepository;',
            'private final ScheduleExecutor scheduleExecutor;',
        ]),
        'operations': '\n'.join([
            '// 스케줄 전체 목록 조회',
            'public List<ScheduleDto.Response> findAll();',
            '',
            '// 스케줄 단건 조회',
            'public ScheduleDto.Response findById(Long scheduleId);',
            '',
            '// Agent별 스케줄 조회',
            'public List<ScheduleDto.Response> findByAgentId(Long agentId);',
            '',
            '// 스케줄 생성',
            'public ScheduleDto.Response create(ScheduleDto.CreateRequest request);',
            '',
            '// 스케줄 수정',
            'public ScheduleDto.Response update(Long scheduleId, ScheduleDto.UpdateRequest request);',
            '',
            '// 스케줄 활성화/비활성화 토글',
            'public ScheduleDto.Response toggle(Long scheduleId);',
            '',
            '// 스케줄 삭제',
            'public void delete(Long scheduleId);',
        ]),
    }]))

    sched_classes.append(('Request', [{
        'class_name': 'ScheduleRequest',
        'class_type': 'Class <<Request>>',
        'overview': make_overview(
            'Schedule Request 정보를 정의한다.',
            '스케줄 생성/수정 요청에 필요한 Agent ID, Cron 표현식, 활성화 여부, 실행 옵션 등의 파라미터를 정의한다.'
        ),
        'attributes': '\n'.join([
            'private Long agentId;               // Agent ID',
            'private String cronExpression;      // Cron 표현식',
            'private Boolean isEnabled;          // 활성화 여부',
            'private String executionOptions;    // 조건실행 옵션 JSON',
        ]),
        'operations': None,
    }]))

    sched_classes.append(('Response', [{
        'class_name': 'ScheduleResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'Schedule Response 정보를 정의한다.',
            '스케줄 조회 응답에 포함되는 스케줄ID, Agent 정보, Cron 표현식, 활성화 여부 등을 정의한다.'
        ),
        'attributes': '\n'.join([
            'private Long scheduleId;            // 스케줄 ID',
            'private Long agentId;               // Agent ID',
            'private String agentCode;           // Agent 코드',
            'private String agentName;           // Agent 명',
            'private String cronExpression;      // Cron 표현식',
            'private Boolean isEnabled;          // 활성화 여부',
            'private String executionOptions;    // 조건실행 옵션 JSON',
            'private LocalDateTime createdAt;    // 생성 시각',
        ]),
        'operations': None,
    }]))

    sched_classes.append(('Entity', [{
        'class_name': 'Schedule',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'Schedule 엔티티를 정의한다.',
            '테이블: schedule. 스케줄 설정 정보를 관리하며, Agent 연관, Cron 표현식, 활성화 여부, 조건실행 옵션(JSON)을 포함한다.'
        ),
        'attributes': '\n'.join([
            '@Id @GeneratedValue(strategy = GenerationType.IDENTITY)',
            '@Column(name = "schedule_id")',
            'private Long scheduleId;                  // PK',
            '',
            '@ManyToOne(fetch = FetchType.LAZY)',
            '@JoinColumn(name = "agent_id", nullable = false)',
            'private Agent agent;                      // Agent (FK)',
            '',
            '@Column(name = "cron_expression", nullable = false)',
            'private String cronExpression;            // Cron 표현식',
            '',
            '@Column(name = "is_enabled")',
            'private Boolean isEnabled;                // 활성화 여부',
            '',
            '@Column(name = "execution_options", columnDefinition = "TEXT")',
            'private String executionOptions;          // 조건실행 옵션 JSON',
            '',
            '@CreationTimestamp',
            '@Column(name = "created_at", updatable = false)',
            'private LocalDateTime createdAt;          // 생성 시각',
        ]),
        'operations': None,
    }]))

    sched_classes.append(('Repository', [{
        'class_name': 'ScheduleRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'ScheduleRepository를 정의한다.',
            'Schedule 엔티티에 대한 JPA Repository로, Agent별 조회, 활성 스케줄 조회, 중복 검사 기능을 제공한다.'
        ),
        'attributes': 'extends JpaRepository<Schedule, Long>',
        'operations': '\n'.join([
            '// Agent별 스케줄 조회',
            'List<Schedule> findByAgentId(Long agentId);',
            '',
            '// 활성화된 스케줄 조회',
            'List<Schedule> findByIsEnabledTrue();',
            '',
            '// 활성 스케줄 + Agent 조회 (FETCH JOIN)',
            '@Query("SELECT s FROM Schedule s JOIN FETCH s.agent WHERE s.isEnabled = true")',
            'List<Schedule> findEnabledSchedulesWithAgent();',
            '',
            '// Agent + Cron 중복 검사',
            'Optional<Schedule> findByAgentIdAndCronExpression(Long agentId, String cronExpression);',
        ]),
    }]))

    groups.append(('Schedule', sched_classes))

    # ================================================================
    # GROUP 7: Zone
    # ================================================================
    zone_classes = []

    zone_classes.append(('Entity', [{
        'class_name': 'ZoneConfig',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'ZoneConfig 엔티티를 정의한다.',
            '테이블: zone_config. 네트워크 존 설정 정보를 관리하며, 존명/약어/프록시 URL/설명/활성화 여부를 포함한다.'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@Column(name = "zone")',
            'private String zone;                      // 존 이름 (PK)',
            '',
            '@Column(name = "short_code", nullable = false)',
            'private String shortCode;                 // 존 약어 (E/D/IC/IS)',
            '',
            '@Column(name = "proxy_agent_url", nullable = false)',
            'private String proxyAgentUrl;             // 프록시 에이전트 URL',
            '',
            '@Column(name = "description")',
            'private String description;               // 설명',
            '',
            '@Column(name = "is_active")',
            'private Boolean isActive;                 // 활성화 여부',
            '',
            '@CreationTimestamp',
            '@Column(name = "created_at", updatable = false)',
            'private LocalDateTime createdAt;          // 생성 시각',
            '',
            '@UpdateTimestamp',
            '@Column(name = "updated_at")',
            'private LocalDateTime updatedAt;          // 수정 시각',
        ]),
        'operations': None,
    }]))

    zone_classes.append(('Repository', [{
        'class_name': 'ZoneConfigRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'ZoneConfigRepository를 정의한다.',
            'ZoneConfig 엔티티에 대한 JPA Repository로, 활성 존 조회, 존 약어 조회 기능을 제공한다.'
        ),
        'attributes': 'extends JpaRepository<ZoneConfig, String>',
        'operations': '\n'.join([
            '// 활성화된 존 조회',
            'Optional<ZoneConfig> findByZoneAndIsActiveTrue(String zone);',
            '',
            '// 존 약어 조회',
            'default String findShortCodeByZone(String zone);',
        ]),
    }]))

    groups.append(('Zone', zone_classes))

    return groups


# ============================================================
# Package mapping per group
# ============================================================
PACKAGE_MAP = {
    'Agent': 'kr.go.gims.orchestrator.domain.agent',
    'Callback': 'kr.go.gims.orchestrator.domain.callback',
    'Chain': 'kr.go.gims.orchestrator.domain.chain',
    'Datasource': 'kr.go.gims.orchestrator.domain.datasource',
    'Execution': 'kr.go.gims.orchestrator.domain.execution',
    'Schedule': 'kr.go.gims.orchestrator.domain.schedule',
    'Zone': 'kr.go.gims.orchestrator.domain.zone',
}


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(9)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    doc.add_heading('sync-orchestrator 클래스명세서', level=0)

    groups = build_all_classes()

    for group_idx, (group_name, type_list) in enumerate(groups, start=1):
        # Group heading (level 1)
        doc.add_heading(f'{group_idx}    {group_name}', level=1)

        sub_idx = 0
        for type_label, class_list in type_list:
            sub_idx += 1
            # Type heading (level 2)
            doc.add_heading(f'{group_idx}.{sub_idx}    {type_label}', level=2)

            for cls_idx, cls in enumerate(class_list, start=1):
                # Class heading (level 3)
                doc.add_heading(
                    f'{group_idx}.{sub_idx}.{cls_idx}    {cls["class_name"]}',
                    level=3
                )

                create_class_table(
                    doc,
                    package_name=PACKAGE_MAP[group_name],
                    class_id='',
                    class_name=cls['class_name'],
                    screen_id='',
                    class_type=cls['class_type'],
                    class_overview=cls['overview'],
                    attributes_text=cls['attributes'],
                    operations_text=cls.get('operations'),
                )

    # Save
    output_dir = r'D:\dev\claude\GIMS\orchestrator_v2\dev_plan\2026_03\19\output'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(
        output_dir, 'sync-orchestrator_NGIS_D23_클래스명세서v0.1.docx'
    )
    doc.save(output_path)
    print(f'클래스명세서 생성 완료: {output_path}')
    print(f'파일 크기: {os.path.getsize(output_path)} bytes')

    # Count classes for verification
    total = sum(len(cl) for _, tl in groups for _, cl in tl)
    print(f'총 클래스 수: {total}개')


if __name__ == '__main__':
    main()
