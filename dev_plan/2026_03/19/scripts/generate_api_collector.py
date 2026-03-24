# -*- coding: utf-8 -*-
"""
infolink-api-collector NGIS D23 클래스명세서 생성 스크립트
- 3개 도메인 그룹, 18개 클래스
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

YELLOW_SHADING = 'FFFF00'
COL_WIDTHS_DXA = [1800, 2701, 1529, 2902]
ROW_HEIGHTS_WITH_OPS = [311, 321, 297, 297, 1300, 261, 275, 261, 258]
ROW_HEIGHTS_NO_OPS = [311, 321, 297, 297, 1300, 261, 275]


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _make_run(p, text, bold=False, font_size=9):
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
    rPr.append(rFonts)
    return run


def set_cell_text(cell, text, bold=False, font_size=9, alignment=None):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _make_run(p, text, bold=bold, font_size=font_size)


def set_cell_multiline(cell, lines_list, bold=False, font_size=9):
    for p in cell.paragraphs:
        p.clear()
    for i, line in enumerate(lines_list):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _make_run(p, line, bold=bold, font_size=font_size)


def merge_cells_horizontal(table, row_idx, start_col, end_col):
    table.cell(row_idx, start_col).merge(table.cell(row_idx, end_col))


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))


def set_table_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing = tblPr.find(qn('w:tblLayout'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(layout)


def set_table_indent(table, indent_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    indent_xml = f'<w:tblInd {nsdecls("w")} w:w="{indent_dxa}" w:type="dxa"/>'
    existing = tblPr.find(qn('w:tblInd'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(indent_xml))


def set_row_height(row, height_dxa):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{height_dxa}" w:hRule="atLeast"/>')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def set_column_widths(table, widths_dxa):
    tbl = table._tbl
    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in widths_dxa:
        grid_xml += f'<w:gridCol w:w="{w}"/>'
    grid_xml += '</w:tblGrid>'
    tblPr = tbl.tblPr
    tblPr.addnext(parse_xml(grid_xml))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_dxa):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                tcPr.append(parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{widths_dxa[i]}" w:type="dxa"/>'
                ))


def set_cell_border_nil(cell, sides):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_parts = []
    for side in sides:
        borders_parts.append(f'<w:{side} w:val="nil"/>')
    borders_xml = f'<w:tcBorders {nsdecls("w")}>{"".join(borders_parts)}</w:tcBorders>'
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(borders_xml))


def create_class_table(doc, package_name, class_id, class_name, screen_id,
                       class_type, class_overview, attributes_text,
                       operations_text=None):
    """Create a single class specification table."""
    has_ops = operations_text is not None
    num_rows = 9 if has_ops else 7

    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table)
    set_table_layout_fixed(table)
    set_table_indent(table, 125)
    set_column_widths(table, COL_WIDTHS_DXA)

    heights = ROW_HEIGHTS_WITH_OPS if has_ops else ROW_HEIGHTS_NO_OPS
    for i, row in enumerate(table.rows):
        if i < len(heights):
            set_row_height(row, heights[i])

    # Row 0: Header
    merge_cells_horizontal(table, 0, 0, 3)
    cell = table.cell(0, 0)
    set_cell_text(cell, '클래스 명세서[설계]', bold=True, font_size=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 1: 패키지명 / 클래스ID
    set_cell_text(table.cell(1, 0), '패키지명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(1, 0), YELLOW_SHADING)
    set_cell_text(table.cell(1, 1), package_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 2), '클래스ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(1, 2), YELLOW_SHADING)
    set_cell_text(table.cell(1, 3), class_id, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 2: 클래스명 / 화면ID
    set_cell_text(table.cell(2, 0), '클래스명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(2, 0), YELLOW_SHADING)
    set_cell_text(table.cell(2, 1), class_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 2), '화면ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(2, 2), YELLOW_SHADING)
    set_cell_text(table.cell(2, 3), screen_id, font_size=9)

    # Row 3: 클래스 타입
    set_cell_text(table.cell(3, 0), '클래스 타입', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(3, 0), YELLOW_SHADING)
    set_cell_text(table.cell(3, 1), class_type, font_size=9)
    set_cell_text(table.cell(3, 2), '', font_size=9)
    set_cell_text(table.cell(3, 3), '', font_size=9)
    set_cell_border_nil(table.cell(3, 1), ['right'])
    set_cell_border_nil(table.cell(3, 2), ['left', 'right'])
    set_cell_border_nil(table.cell(3, 3), ['left'])

    # Row 4: 클래스 개요
    cell0 = table.cell(4, 0)
    set_cell_shading(cell0, YELLOW_SHADING)
    set_cell_multiline(cell0, ['', '', '클래스 개요'], bold=True, font_size=9)
    for p in cell0.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    merge_cells_horizontal(table, 4, 1, 3)
    overview_lines = class_overview.split('\n')
    set_cell_multiline(table.cell(4, 1), overview_lines, font_size=9)

    # Row 5: 속성(Attributes) header
    merge_cells_horizontal(table, 5, 0, 3)
    cell = table.cell(5, 0)
    set_cell_text(cell, '속성(Attributes)', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 6: attributes content
    merge_cells_horizontal(table, 6, 0, 3)
    attr_lines = attributes_text.split('\n')
    set_cell_multiline(table.cell(6, 0), attr_lines, font_size=9)

    if has_ops:
        # Row 7: 오퍼레이션(Operations) header
        merge_cells_horizontal(table, 7, 0, 3)
        cell = table.cell(7, 0)
        set_cell_text(cell, '오퍼레이션(Operations)', bold=True, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, YELLOW_SHADING)

        # Row 8: operations content
        merge_cells_horizontal(table, 8, 0, 3)
        ops_lines = operations_text.split('\n')
        set_cell_multiline(table.cell(8, 0), ops_lines, font_size=9)

    doc.add_paragraph()  # spacing


# ============================================================
# Overview helper
# ============================================================
def make_overview(line1, line2):
    return '\n'.join([
        '개요',
        f'-    {line1}',
        '',
        '상세내용',
        f'-    {line2}',
    ])


# ============================================================
# PACKAGE MAP
# ============================================================
PKG = 'kr.go.gims.collector'

PACKAGE_MAP = {
    'ApiEndpointController': f'{PKG}.controller',
    'ApiScheduleController': f'{PKG}.controller',
    'ApiHistoryController': f'{PKG}.controller',
    'ApiEndpointService': f'{PKG}.service',
    'ApiExecutionService': f'{PKG}.service',
    'ApiScheduleService': f'{PKG}.service',
    'ApiTestService': f'{PKG}.service',
    'ApiEndpointRequest': f'{PKG}.dto',
    'ApiEndpointResponse': f'{PKG}.dto',
    'ApiScheduleRequest': f'{PKG}.dto',
    'ApiScheduleResponse': f'{PKG}.dto',
    'ApiExecutionHistoryResponse': f'{PKG}.dto',
    'TestCallRequest': f'{PKG}.dto',
    'TestCallResponse': f'{PKG}.dto',
    'ApiEndpoint': f'{PKG}.entity',
    'ApiParam': f'{PKG}.entity',
    'ApiFieldMapping': f'{PKG}.entity',
    'ApiExecutionHistory': f'{PKG}.entity',
    'ApiSchedule': f'{PKG}.entity',
    'ApiEndpointRepository': f'{PKG}.repository',
    'ApiParamRepository': f'{PKG}.repository',
    'ApiFieldMappingRepository': f'{PKG}.repository',
    'ApiExecutionHistoryRepository': f'{PKG}.repository',
    'ApiScheduleRepository': f'{PKG}.repository',
}


# ============================================================
# DATA DEFINITIONS
# ============================================================
def build_all_classes():
    """Return list of (group_name, [(type_label, [class_data])]) tuples."""
    groups = []

    # ================================================================
    # GROUP 1: ApiEndpoint
    # ================================================================
    endpoint_classes = []

    # --- Controller ---
    endpoint_classes.append(('Controller', [{
        'class_name': 'ApiEndpointController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'ApiEndpointController를 정의한다.',
            'API 엔드포인트 CRUD, 파라미터 저장, 테스트 호출, 인라인 테스트, 수동 실행, 필드 매핑 CRUD 기능을 제공한다.'
        ),
        'attributes': '\n'.join([
            'private final ApiEndpointService endpointService;',
            'private final ApiTestService testService;',
            'private final ApiExecutionService executionService;',
        ]),
        'operations': '\n'.join([
            '// API 엔드포인트 전체 목록 조회',
            '@GetMapping',
            'public List<ListResponse> getList();',
            '',
            '// API 엔드포인트 단건 상세 조회',
            '@GetMapping("/{id}")',
            'public DetailResponse getDetail(@PathVariable Long id);',
            '',
            '// API 엔드포인트 신규 등록',
            '@PostMapping',
            'public ResponseEntity<DetailResponse> create(@Valid @RequestBody CreateRequest request);',
            '',
            '// API 엔드포인트 수정',
            '@PutMapping("/{id}")',
            'public DetailResponse update(@PathVariable Long id, @Valid @RequestBody UpdateRequest request);',
            '',
            '// API 엔드포인트 삭제',
            '@DeleteMapping("/{id}")',
            'public ResponseEntity<Void> delete(@PathVariable Long id);',
            '',
            '// 파라미터 일괄 저장',
            '@PutMapping("/{id}/params")',
            'public List<ParamResponse> saveParams(@PathVariable Long id, @RequestBody List<ParamRequest> requests);',
            '',
            '// 저장된 엔드포인트 테스트 호출',
            '@PostMapping("/{id}/test")',
            'public TestCallDto.Response testCall(@PathVariable Long id, @RequestBody(required = false) TestCallDto.Request request);',
            '',
            '// 저장 없이 인라인 테스트 (등록 전 검증용)',
            '@PostMapping("/test-inline")',
            'public TestCallDto.Response testCallInline(@RequestBody TestCallDto.InlineRequest request);',
            '',
            '// API 수동 실행',
            '@PostMapping("/{id}/run")',
            'public ApiExecutionHistoryDto.Response run(@PathVariable Long id);',
            '',
            '// 필드 매핑 목록 조회',
            '@GetMapping("/{id}/mappings")',
            'public List<FieldMappingResponse> getMappings(@PathVariable Long id);',
            '',
            '// 필드 매핑 일괄 저장',
            '@PutMapping("/{id}/mappings")',
            'public List<FieldMappingResponse> saveMappings(@PathVariable Long id, @RequestBody List<FieldMappingRequest> requests);',
        ]),
    }]))

    # --- Service ---
    endpoint_classes.append(('Service', [{
        'class_name': 'ApiEndpointService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'ApiEndpointService를 정의한다.',
            'API 엔드포인트 CRUD, 파라미터 일괄 저장, 필드 매핑 일괄 저장 비즈니스 로직을 처리한다.'
        ),
        'attributes': '\n'.join([
            'private final ApiEndpointRepository endpointRepository;',
            'private final ApiParamRepository paramRepository;',
            'private final ApiFieldMappingRepository mappingRepository;',
            'private final ApiScheduleRepository scheduleRepository;',
            'private final ApiExecutionHistoryRepository historyRepository;',
        ]),
        'operations': '\n'.join([
            '// API 엔드포인트 전체 목록 조회',
            '@Transactional(readOnly = true)',
            'public List<ListResponse> getList();',
            '',
            '// API 엔드포인트 단건 상세 조회',
            '@Transactional(readOnly = true)',
            'public DetailResponse getDetail(Long id);',
            '',
            '// API 엔드포인트 신규 등록',
            '@Transactional',
            'public DetailResponse create(CreateRequest request);',
            '',
            '// API 엔드포인트 수정',
            '@Transactional',
            'public DetailResponse update(Long id, UpdateRequest request);',
            '',
            '// API 엔드포인트 삭제 (연관 이력/스케줄 포함)',
            '@Transactional',
            'public void delete(Long id);',
            '',
            '// 파라미터 일괄 저장 (기존 삭제 후 재생성)',
            '@Transactional',
            'public List<ParamResponse> saveParams(Long endpointId, List<ParamRequest> requests);',
            '',
            '// 필드 매핑 일괄 저장 (기존 삭제 후 재생성)',
            '@Transactional',
            'public List<FieldMappingResponse> saveMappings(Long endpointId, List<FieldMappingRequest> requests);',
        ]),
    }]))

    # --- Request ---
    endpoint_classes.append(('Request', [{
        'class_name': 'ApiEndpointRequest',
        'class_type': 'Class <<Request>>',
        'overview': make_overview(
            'ApiEndpointRequest를 정의한다.',
            'API 엔드포인트 생성/수정, 파라미터 저장, 필드 매핑 저장 요청 DTO를 통합한다. (CreateRequest, UpdateRequest, ParamRequest, FieldMappingRequest)'
        ),
        'attributes': '\n'.join([
            '// --- CreateRequest ---',
            'private String apiName; // API 표시명',
            'private String url; // 호출 URL',
            'private String httpMethod; // HTTP 메서드',
            'private String contentType; // Content-Type',
            'private String headers; // 추가 헤더 JSON',
            'private ApiEndpoint.AuthType authType; // 인증 유형',
            'private String authConfig; // 인증 설정 JSON',
            'private String description; // 설명',
            'private ApiEndpoint.Zone zone; // 배포 존',
            '',
            '// --- UpdateRequest 추가 필드 ---',
            'private String dataRootPath; // 응답 데이터 배열 경로',
            'private String targetDatasourceId; // 적재 대상 datasource ID',
            'private String targetTableName; // 적재 대상 테이블명',
            'private Boolean upsertEnabled; // UPSERT 사용 여부',
            'private Boolean isActive; // 활성화 여부',
            '',
            '// --- ParamRequest ---',
            'private String paramName; // 파라미터명',
            'private ApiParam.ParamType paramType; // 파라미터 위치',
            'private ApiParam.ValueType valueType; // 값 유형',
            'private String staticValue; // 고정값',
            'private ApiParam.DynamicType dynamicType; // 동적 유형',
            'private String dynamicFormat; // 날짜 포맷',
            'private Integer dynamicOffset; // 오프셋',
            'private Integer displayOrder; // 표시 순서',
            '',
            '// --- FieldMappingRequest ---',
            'private String sourceFieldPath; // 응답 필드 경로',
            'private String targetColumnName; // DB 컬럼명',
            'private Boolean isPk; // PK 여부',
            'private ApiFieldMapping.TransformType transformType; // 변환 유형',
            'private String transformConfig; // 변환 설정 JSON',
        ]),
    }]))

    # --- Response ---
    endpoint_classes.append(('Response', [{
        'class_name': 'ApiEndpointResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'ApiEndpointResponse를 정의한다.',
            'API 엔드포인트 목록/상세, 파라미터, 필드 매핑 응답 DTO를 통합한다. (ListResponse, DetailResponse, ParamResponse, FieldMappingResponse)'
        ),
        'attributes': '\n'.join([
            '// --- ListResponse ---',
            'private Long id; // PK',
            'private String apiName; // API 표시명',
            'private String url; // 호출 URL',
            'private String httpMethod; // HTTP 메서드',
            'private ApiEndpoint.AuthType authType; // 인증 유형',
            'private String targetTableName; // 적재 대상 테이블명',
            'private Boolean isActive; // 활성화 여부',
            'private ApiEndpoint.Zone zone; // 배포 존',
            'private boolean hasMappings; // 매핑 존재 여부',
            'private LocalDateTime createdAt; // 생성 시각',
            'private LocalDateTime updatedAt; // 수정 시각',
            '',
            '// --- DetailResponse 추가 필드 ---',
            'private String contentType; // Content-Type',
            'private String headers; // 추가 헤더 JSON',
            'private String authConfig; // 인증 설정 JSON',
            'private String dataRootPath; // 응답 데이터 배열 경로',
            'private String targetDatasourceId; // 적재 대상 datasource ID',
            'private Boolean upsertEnabled; // UPSERT 사용 여부',
            'private String description; // 설명',
            'private List<ParamResponse> params; // 파라미터 목록',
            'private List<FieldMappingResponse> fieldMappings; // 필드 매핑 목록',
            '',
            '// --- ParamResponse ---',
            'private Long id; // PK',
            'private String paramName; // 파라미터명',
            'private ApiParam.ParamType paramType; // 파라미터 위치',
            'private ApiParam.ValueType valueType; // 값 유형',
            'private String staticValue; // 고정값',
            'private ApiParam.DynamicType dynamicType; // 동적 유형',
            'private String dynamicFormat; // 날짜 포맷',
            'private Integer dynamicOffset; // 오프셋',
            'private String description; // 설명',
            'private Integer displayOrder; // 표시 순서',
            '',
            '// --- FieldMappingResponse ---',
            'private Long id; // PK',
            'private String sourceFieldPath; // 응답 필드 경로',
            'private String targetColumnName; // DB 컬럼명',
            'private Boolean isPk; // PK 여부',
            'private ApiFieldMapping.TransformType transformType; // 변환 유형',
            'private String transformConfig; // 변환 설정 JSON',
            'private Integer displayOrder; // 표시 순서',
        ]),
    }]))

    # --- Entity ---
    endpoint_entities = []

    # ApiEndpoint entity
    endpoint_entities.append({
        'class_name': 'ApiEndpoint',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'ApiEndpoint를 정의한다.',
            '외부 API 엔드포인트 정의 엔티티이다. 테이블명: api_endpoint'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            'private Long id;',
            '',
            '@Column(name = "api_name", length = 100, nullable = false)',
            'private String apiName;',
            '',
            '@Column(name = "url", length = 500, nullable = false)',
            'private String url;',
            '',
            '@Column(name = "http_method", length = 10, nullable = false)',
            'private String httpMethod;',
            '',
            '@Column(name = "content_type", length = 50)',
            'private String contentType;',
            '',
            '@Column(name = "headers", columnDefinition = "TEXT")',
            'private String headers;',
            '',
            '@Enumerated(EnumType.STRING)',
            '@Column(name = "auth_type", length = 20, nullable = false)',
            'private AuthType authType;',
            '',
            '@Column(name = "auth_config", columnDefinition = "TEXT")',
            'private String authConfig;',
            '',
            '@Column(name = "data_root_path", length = 200)',
            'private String dataRootPath;',
            '',
            '@Column(name = "target_datasource_id", length = 50)',
            'private String targetDatasourceId;',
            '',
            '@Column(name = "target_table_name", length = 100)',
            'private String targetTableName;',
            '',
            '@Column(name = "upsert_enabled")',
            'private Boolean upsertEnabled = false;',
            '',
            '@Column(name = "description", length = 500)',
            'private String description;',
            '',
            '@Column(name = "is_active")',
            'private Boolean isActive = true;',
            '',
            '@Enumerated(EnumType.STRING)',
            '@Column(name = "zone", length = 20, nullable = false)',
            'private Zone zone;',
            '',
            '@CreationTimestamp',
            '@Column(name = "created_at", updatable = false)',
            'private LocalDateTime createdAt;',
            '',
            '@UpdateTimestamp',
            '@Column(name = "updated_at")',
            'private LocalDateTime updatedAt;',
            '',
            '@OneToMany(mappedBy = "apiEndpoint", cascade = CascadeType.ALL, orphanRemoval = true)',
            '@OrderBy("displayOrder ASC")',
            'private List<ApiParam> params = new ArrayList<>();',
            '',
            '@OneToMany(mappedBy = "apiEndpoint", cascade = CascadeType.ALL, orphanRemoval = true)',
            '@OrderBy("displayOrder ASC")',
            'private List<ApiFieldMapping> fieldMappings = new ArrayList<>();',
        ]),
    })

    # ApiParam entity
    endpoint_entities.append({
        'class_name': 'ApiParam',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'ApiParam을 정의한다.',
            'API 호출 파라미터 엔티티이다. 테이블명: api_param'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            'private Long id;',
            '',
            '@ManyToOne(fetch = FetchType.LAZY)',
            '@JoinColumn(name = "api_endpoint_id", nullable = false)',
            'private ApiEndpoint apiEndpoint;',
            '',
            '@Column(name = "param_name", length = 100, nullable = false)',
            'private String paramName;',
            '',
            '@Enumerated(EnumType.STRING)',
            '@Column(name = "param_type", length = 20, nullable = false)',
            'private ParamType paramType;',
            '',
            '@Enumerated(EnumType.STRING)',
            '@Column(name = "value_type", length = 20, nullable = false)',
            'private ValueType valueType;',
            '',
            '@Column(name = "static_value", length = 500)',
            'private String staticValue;',
            '',
            '@Enumerated(EnumType.STRING)',
            '@Column(name = "dynamic_type", length = 20)',
            'private DynamicType dynamicType;',
            '',
            '@Column(name = "dynamic_format", length = 50)',
            'private String dynamicFormat;',
            '',
            '@Column(name = "dynamic_offset")',
            'private Integer dynamicOffset;',
            '',
            '@Column(name = "description", length = 200)',
            'private String description;',
            '',
            '@Column(name = "display_order")',
            'private Integer displayOrder = 0;',
        ]),
    })

    # ApiFieldMapping entity
    endpoint_entities.append({
        'class_name': 'ApiFieldMapping',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'ApiFieldMapping을 정의한다.',
            'API 응답 필드와 DB 컬럼 간 매핑 엔티티이다. 테이블명: api_field_mapping'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            'private Long id;',
            '',
            '@ManyToOne(fetch = FetchType.LAZY)',
            '@JoinColumn(name = "api_endpoint_id", nullable = false)',
            'private ApiEndpoint apiEndpoint;',
            '',
            '@Column(name = "source_field_path", length = 200, nullable = false)',
            'private String sourceFieldPath;',
            '',
            '@Column(name = "target_column_name", length = 100, nullable = false)',
            'private String targetColumnName;',
            '',
            '@Column(name = "is_pk")',
            'private Boolean isPk = false;',
            '',
            '@Enumerated(EnumType.STRING)',
            '@Column(name = "transform_type", length = 20, nullable = false)',
            'private TransformType transformType = TransformType.NONE;',
            '',
            '@Column(name = "transform_config", columnDefinition = "TEXT")',
            'private String transformConfig;',
            '',
            '@Column(name = "display_order")',
            'private Integer displayOrder = 0;',
        ]),
    })

    endpoint_classes.append(('Entity', endpoint_entities))

    # --- Repository ---
    endpoint_repos = []

    endpoint_repos.append({
        'class_name': 'ApiEndpointRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'ApiEndpointRepository를 정의한다.',
            'ApiEndpoint 엔티티에 대한 JPA Repository이다.'
        ),
        'attributes': 'extends JpaRepository<ApiEndpoint, Long>',
    })

    endpoint_repos.append({
        'class_name': 'ApiParamRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'ApiParamRepository를 정의한다.',
            'ApiParam 엔티티에 대한 JPA Repository이다.'
        ),
        'attributes': 'extends JpaRepository<ApiParam, Long>',
        'operations': '\n'.join([
            '// 엔드포인트 ID 기준 파라미터 목록 조회 (표시 순서 오름차순)',
            'List<ApiParam> findByApiEndpointIdOrderByDisplayOrderAsc(Long endpointId);',
            '',
            '// 엔드포인트 ID 기준 파라미터 전체 삭제',
            'void deleteByApiEndpointId(Long endpointId);',
        ]),
    })

    endpoint_repos.append({
        'class_name': 'ApiFieldMappingRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'ApiFieldMappingRepository를 정의한다.',
            'ApiFieldMapping 엔티티에 대한 JPA Repository이다.'
        ),
        'attributes': 'extends JpaRepository<ApiFieldMapping, Long>',
        'operations': '\n'.join([
            '// 엔드포인트 ID 기준 필드 매핑 목록 조회 (표시 순서 오름차순)',
            'List<ApiFieldMapping> findByApiEndpointIdOrderByDisplayOrderAsc(Long endpointId);',
            '',
            '// 엔드포인트 ID 기준 필드 매핑 전체 삭제',
            'void deleteByApiEndpointId(Long endpointId);',
        ]),
    })

    endpoint_classes.append(('Repository', endpoint_repos))

    groups.append(('ApiEndpoint', endpoint_classes))

    # ================================================================
    # GROUP 2: ApiSchedule
    # ================================================================
    schedule_classes = []

    # --- Controller ---
    schedule_classes.append(('Controller', [{
        'class_name': 'ApiScheduleController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'ApiScheduleController를 정의한다.',
            'API 수집 스케줄 CRUD 및 활성화 토글 기능을 제공한다.'
        ),
        'attributes': 'private final ApiScheduleService scheduleService;',
        'operations': '\n'.join([
            '// 엔드포인트별 스케줄 목록 조회',
            '@GetMapping("/api/endpoints/{endpointId}/schedules")',
            'public List<Response> getSchedules(@PathVariable Long endpointId);',
            '',
            '// 스케줄 신규 등록',
            '@PostMapping("/api/endpoints/{endpointId}/schedules")',
            'public ResponseEntity<Response> create(@PathVariable Long endpointId, @Valid @RequestBody CreateRequest request);',
            '',
            '// 스케줄 수정',
            '@PutMapping("/api/schedules/{id}")',
            'public Response update(@PathVariable Long id, @Valid @RequestBody UpdateRequest request);',
            '',
            '// 스케줄 활성화/비활성화 토글',
            '@PutMapping("/api/schedules/{id}/toggle")',
            'public Response toggle(@PathVariable Long id);',
            '',
            '// 스케줄 삭제',
            '@DeleteMapping("/api/schedules/{id}")',
            'public ResponseEntity<Void> delete(@PathVariable Long id);',
        ]),
    }]))

    # --- Service ---
    schedule_classes.append(('Service', [{
        'class_name': 'ApiScheduleService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'ApiScheduleService를 정의한다.',
            'API 수집 스케줄 CRUD, 활성화 토글 비즈니스 로직을 처리하고 스케줄 실행기를 연동한다.'
        ),
        'attributes': '\n'.join([
            'private final ApiScheduleRepository scheduleRepository;',
            'private final ApiEndpointRepository endpointRepository;',
            'private final ApiScheduleExecutor scheduleExecutor;',
        ]),
        'operations': '\n'.join([
            '// 엔드포인트별 스케줄 목록 조회',
            '@Transactional(readOnly = true)',
            'public List<Response> getSchedules(Long endpointId);',
            '',
            '// 스케줄 신규 등록 및 실행기 등록',
            '@Transactional',
            'public Response create(Long endpointId, CreateRequest request);',
            '',
            '// 스케줄 수정 및 실행기 재등록',
            '@Transactional',
            'public Response update(Long scheduleId, UpdateRequest request);',
            '',
            '// 스케줄 활성화/비활성화 토글',
            '@Transactional',
            'public Response toggle(Long scheduleId);',
            '',
            '// 스케줄 삭제 및 실행기 해제',
            '@Transactional',
            'public void delete(Long scheduleId);',
        ]),
    }]))

    # --- Request ---
    schedule_classes.append(('Request', [{
        'class_name': 'ApiScheduleRequest',
        'class_type': 'Class <<Request>>',
        'overview': make_overview(
            'ApiScheduleRequest를 정의한다.',
            'API 수집 스케줄 생성/수정 요청 DTO를 통합한다. (CreateRequest, UpdateRequest)'
        ),
        'attributes': '\n'.join([
            '// --- CreateRequest ---',
            'private String cronExpression; // Cron 표현식',
            '',
            '// --- UpdateRequest ---',
            'private String cronExpression; // Cron 표현식',
        ]),
    }]))

    # --- Response ---
    schedule_classes.append(('Response', [{
        'class_name': 'ApiScheduleResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'ApiScheduleResponse를 정의한다.',
            'API 수집 스케줄 응답 DTO이다.'
        ),
        'attributes': '\n'.join([
            'private Long id; // PK',
            'private Long apiEndpointId; // 엔드포인트 ID',
            'private String cronExpression; // Cron 표현식',
            'private Boolean isEnabled; // 활성화 여부',
            'private LocalDateTime createdAt; // 생성 시각',
        ]),
    }]))

    # --- Entity ---
    schedule_classes.append(('Entity', [{
        'class_name': 'ApiSchedule',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'ApiSchedule을 정의한다.',
            'API 수집 스케줄 엔티티이다. 테이블명: api_schedule'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            'private Long id;',
            '',
            '@ManyToOne(fetch = FetchType.LAZY)',
            '@JoinColumn(name = "api_endpoint_id", nullable = false)',
            'private ApiEndpoint apiEndpoint;',
            '',
            '@Column(name = "cron_expression", length = 50, nullable = false)',
            'private String cronExpression;',
            '',
            '@Column(name = "is_enabled")',
            'private Boolean isEnabled = true;',
            '',
            '@CreationTimestamp',
            '@Column(name = "created_at", updatable = false)',
            'private LocalDateTime createdAt;',
        ]),
    }]))

    # --- Repository ---
    schedule_classes.append(('Repository', [{
        'class_name': 'ApiScheduleRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'ApiScheduleRepository를 정의한다.',
            'ApiSchedule 엔티티에 대한 JPA Repository이다.'
        ),
        'attributes': 'extends JpaRepository<ApiSchedule, Long>',
        'operations': '\n'.join([
            '// 엔드포인트 ID 기준 스케줄 목록 조회',
            'List<ApiSchedule> findByApiEndpointId(Long endpointId);',
            '',
            '// 활성화된 스케줄 전체 조회',
            'List<ApiSchedule> findByIsEnabledTrue();',
            '',
            '// 엔드포인트 ID 기준 스케줄 전체 삭제',
            'void deleteByApiEndpointId(Long endpointId);',
        ]),
    }]))

    groups.append(('ApiSchedule', schedule_classes))

    # ================================================================
    # GROUP 3: ApiHistory
    # ================================================================
    history_classes = []

    # --- Controller ---
    history_classes.append(('Controller', [{
        'class_name': 'ApiHistoryController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'ApiHistoryController를 정의한다.',
            'API 수집 실행 이력 페이징 조회 기능을 제공한다.'
        ),
        'attributes': 'private final ApiExecutionHistoryRepository historyRepository;',
        'operations': '\n'.join([
            '// 엔드포인트별 실행 이력 페이징 조회',
            '@GetMapping',
            'public Page<ApiExecutionHistoryDto.Response> getHistory(',
            '        @PathVariable Long endpointId,',
            '        @RequestParam(defaultValue = "0") int page,',
            '        @RequestParam(defaultValue = "20") int size);',
        ]),
    }]))

    # --- Service ---
    history_services = []

    history_services.append({
        'class_name': 'ApiExecutionService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'ApiExecutionService를 정의한다.',
            '범용 API 실행 엔진이다. 설정 로드, API 호출, 응답 파싱, 매핑 규칙 적용, JDBC batch INSERT/UPSERT, 이력 기록을 수행한다.'
        ),
        'attributes': '\n'.join([
            'private final ApiEndpointRepository endpointRepository;',
            'private final ApiExecutionHistoryRepository historyRepository;',
            'private final ApiCallService callService;',
            'private final ResponseParser responseParser;',
            'private final DataTransformer dataTransformer;',
            'private final DataSource dataSource;',
            'private final DynamicDataSourceService dynamicDataSourceService;',
        ]),
        'operations': '\n'.join([
            '// API 실행 (호출 → 파싱 → 매핑 → DB 적재 → 이력 기록)',
            '@Transactional',
            'public ApiExecutionHistoryDto.Response run(Long endpointId, ApiExecutionHistory.TriggeredBy triggeredBy);',
            '',
            '// 실행 이력 종료 처리 (종료 시각, 소요 시간 기록)',
            'private void finishHistory(ApiExecutionHistory history, LocalDateTime startedAt);',
            '',
            '// Target DataSource 결정 (외부 DB 또는 자체 DB)',
            'private DataSource resolveTargetDataSource(ApiEndpoint endpoint);',
            '',
            '// INSERT 또는 UPSERT SQL 생성 (PostgreSQL ON CONFLICT)',
            'private String buildSql(String tableName, List<String> columns, List<String> pkColumns, boolean upsert);',
            '',
            '// dot notation으로 중첩 값 추출',
            'private Object getNestedValue(Map<String, Object> record, String path);',
        ]),
    })

    history_services.append({
        'class_name': 'ApiTestService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'ApiTestService를 정의한다.',
            'API 테스트 호출 서비스이다. 저장된 엔드포인트 테스트와 인라인 테스트를 지원한다.'
        ),
        'attributes': '\n'.join([
            'private final ApiEndpointRepository endpointRepository;',
            'private final ApiCallService callService;',
            'private final DynamicParamResolver paramResolver;',
            'private final ResponseParser responseParser;',
        ]),
        'operations': '\n'.join([
            '// 저장된 엔드포인트 테스트 호출 (파라미터 오버라이드 지원)',
            '@Transactional(readOnly = true)',
            'public TestCallDto.Response testCall(Long endpointId, Map<String, String> paramOverrides);',
            '',
            '// 저장 없이 인라인 테스트 (임시 객체 조립 후 호출)',
            'public TestCallDto.Response testCallInline(TestCallDto.InlineRequest req);',
        ]),
    })

    history_classes.append(('Service', history_services))

    # --- Response ---
    history_classes.append(('Response', [{
        'class_name': 'ApiExecutionHistoryResponse',
        'class_type': 'Class <<Response>>',
        'overview': make_overview(
            'ApiExecutionHistoryResponse를 정의한다.',
            'API 수집 실행 이력 응답 DTO이다.'
        ),
        'attributes': '\n'.join([
            'private Long id; // PK',
            'private String executionId; // 실행 UUID',
            'private ApiExecutionHistory.Status status; // 실행 상태',
            'private Integer httpStatusCode; // HTTP 응답 코드',
            'private Integer responseCount; // 파싱된 레코드 수',
            'private Integer insertCount; // 적재 성공 수',
            'private Integer skipCount; // 스킵 수',
            'private String errorMessage; // 에러 메시지',
            'private LocalDateTime startedAt; // 시작 시각',
            'private LocalDateTime finishedAt; // 종료 시각',
            'private Long durationMs; // 소요 시간 (ms)',
            'private ApiExecutionHistory.TriggeredBy triggeredBy; // 트리거 유형',
        ]),
    }]))

    # --- Entity ---
    history_classes.append(('Entity', [{
        'class_name': 'ApiExecutionHistory',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'ApiExecutionHistory를 정의한다.',
            'API 수집 실행 이력 엔티티이다. 테이블명: api_execution_history'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            'private Long id;',
            '',
            '@ManyToOne(fetch = FetchType.LAZY)',
            '@JoinColumn(name = "api_endpoint_id", nullable = false)',
            'private ApiEndpoint apiEndpoint;',
            '',
            '@Column(name = "execution_id", length = 36, nullable = false)',
            'private String executionId;',
            '',
            '@Enumerated(EnumType.STRING)',
            '@Column(name = "status", length = 20, nullable = false)',
            'private Status status;',
            '',
            '@Column(name = "http_status_code")',
            'private Integer httpStatusCode;',
            '',
            '@Column(name = "response_count")',
            'private Integer responseCount;',
            '',
            '@Column(name = "insert_count")',
            'private Integer insertCount;',
            '',
            '@Column(name = "skip_count")',
            'private Integer skipCount;',
            '',
            '@Column(name = "error_message", columnDefinition = "TEXT")',
            'private String errorMessage;',
            '',
            '@Column(name = "started_at")',
            'private LocalDateTime startedAt;',
            '',
            '@Column(name = "finished_at")',
            'private LocalDateTime finishedAt;',
            '',
            '@Column(name = "duration_ms")',
            'private Long durationMs;',
            '',
            '@Enumerated(EnumType.STRING)',
            '@Column(name = "triggered_by", length = 20, nullable = false)',
            'private TriggeredBy triggeredBy;',
        ]),
    }]))

    # --- Repository ---
    history_classes.append(('Repository', [{
        'class_name': 'ApiExecutionHistoryRepository',
        'class_type': 'Interface <<Repository>>',
        'overview': make_overview(
            'ApiExecutionHistoryRepository를 정의한다.',
            'ApiExecutionHistory 엔티티에 대한 JPA Repository이다.'
        ),
        'attributes': 'extends JpaRepository<ApiExecutionHistory, Long>',
        'operations': '\n'.join([
            '// 엔드포인트 ID 기준 실행 이력 페이징 조회 (시작 시각 역순)',
            'Page<ApiExecutionHistory> findByApiEndpointIdOrderByStartedAtDesc(Long endpointId, Pageable pageable);',
            '',
            '// 엔드포인트 ID 기준 실행 이력 전체 삭제',
            'void deleteByApiEndpointId(Long endpointId);',
        ]),
    }]))

    groups.append(('ApiHistory', history_classes))

    return groups


# ============================================================
# MAIN
# ============================================================
def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(9)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    doc.add_heading('infolink-api-collector 클래스명세서', level=0)

    groups = build_all_classes()

    for group_idx, (group_name, type_list) in enumerate(groups, start=1):
        # Group heading (level 1)
        doc.add_heading(f'{group_idx}    {group_name}', level=1)

        sub_idx = 0
        for type_label, class_list in type_list:
            sub_idx += 1
            # Type heading (level 2)
            doc.add_heading(f'{group_idx}.{sub_idx}    {type_label}', level=2)

            for cls_idx, cls in enumerate(class_list, start=1):
                # Class heading (level 3)
                doc.add_heading(
                    f'{group_idx}.{sub_idx}.{cls_idx}    {cls["class_name"]}',
                    level=3
                )

                create_class_table(
                    doc,
                    package_name=PACKAGE_MAP[cls['class_name']],
                    class_id='',
                    class_name=cls['class_name'],
                    screen_id='',
                    class_type=cls['class_type'],
                    class_overview=cls['overview'],
                    attributes_text=cls['attributes'],
                    operations_text=cls.get('operations'),
                )

    # Save
    output_dir = r'D:\dev\claude\GIMS\orchestrator_v2\dev_plan\2026_03\19\output'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(
        output_dir, 'infolink-api-collector_NGIS_D23_클래스명세서v0.1.docx'
    )
    doc.save(output_path)
    print(f'클래스명세서 생성 완료: {output_path}')
    print(f'파일 크기: {os.path.getsize(output_path)} bytes')

    # Count classes for verification
    total = sum(len(cl) for _, tl in groups for _, cl in tl)
    print(f'총 클래스 수: {total}개')


if __name__ == '__main__':
    main()
