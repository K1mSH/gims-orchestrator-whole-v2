# -*- coding: utf-8 -*-
"""
sync-agent-bojo NGIS D23 클래스명세서 생성 스크립트
- 2개 도메인 그룹, 14개 클래스
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

YELLOW_SHADING = 'FFFF00'
COL_WIDTHS_DXA = [1800, 2701, 1529, 2902]
ROW_HEIGHTS_WITH_OPS = [311, 321, 297, 297, 1300, 261, 275, 261, 258]
ROW_HEIGHTS_NO_OPS = [311, 321, 297, 297, 1300, 261, 275]


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _make_run(p, text, bold=False, font_size=9):
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
    rPr.append(rFonts)
    return run


def set_cell_text(cell, text, bold=False, font_size=9, alignment=None):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _make_run(p, text, bold=bold, font_size=font_size)


def set_cell_multiline(cell, lines_list, bold=False, font_size=9):
    for p in cell.paragraphs:
        p.clear()
    for i, line in enumerate(lines_list):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _make_run(p, line, bold=bold, font_size=font_size)


def merge_cells_horizontal(table, row_idx, start_col, end_col):
    table.cell(row_idx, start_col).merge(table.cell(row_idx, end_col))


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))


def set_table_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing = tblPr.find(qn('w:tblLayout'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(layout)


def set_table_indent(table, indent_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    indent_xml = f'<w:tblInd {nsdecls("w")} w:w="{indent_dxa}" w:type="dxa"/>'
    existing = tblPr.find(qn('w:tblInd'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(indent_xml))


def set_row_height(row, height_dxa):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{height_dxa}" w:hRule="atLeast"/>')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def set_column_widths(table, widths_dxa):
    tbl = table._tbl
    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in widths_dxa:
        grid_xml += f'<w:gridCol w:w="{w}"/>'
    grid_xml += '</w:tblGrid>'
    tblPr = tbl.tblPr
    tblPr.addnext(parse_xml(grid_xml))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_dxa):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                tcPr.append(parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{widths_dxa[i]}" w:type="dxa"/>'
                ))


def set_cell_border_nil(cell, sides):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_parts = []
    for side in sides:
        borders_parts.append(f'<w:{side} w:val="nil"/>')
    borders_xml = f'<w:tcBorders {nsdecls("w")}>{"".join(borders_parts)}</w:tcBorders>'
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(borders_xml))


def create_class_table(doc, package_name, class_id, class_name, screen_id,
                       class_type, class_overview, attributes_text,
                       operations_text=None):
    """Create a single class specification table."""
    has_ops = operations_text is not None
    num_rows = 9 if has_ops else 7

    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table)
    set_table_layout_fixed(table)
    set_table_indent(table, 125)
    set_column_widths(table, COL_WIDTHS_DXA)

    heights = ROW_HEIGHTS_WITH_OPS if has_ops else ROW_HEIGHTS_NO_OPS
    for i, row in enumerate(table.rows):
        if i < len(heights):
            set_row_height(row, heights[i])

    # Row 0: Header
    merge_cells_horizontal(table, 0, 0, 3)
    cell = table.cell(0, 0)
    set_cell_text(cell, '클래스 명세서[설계]', bold=True, font_size=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 1
    set_cell_text(table.cell(1, 0), '패키지명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(1, 0), YELLOW_SHADING)
    set_cell_text(table.cell(1, 1), package_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 2), '클래스ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(1, 2), YELLOW_SHADING)
    set_cell_text(table.cell(1, 3), class_id, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 2
    set_cell_text(table.cell(2, 0), '클래스명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(2, 0), YELLOW_SHADING)
    set_cell_text(table.cell(2, 1), class_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 2), '화면ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(2, 2), YELLOW_SHADING)
    set_cell_text(table.cell(2, 3), screen_id, font_size=9)

    # Row 3
    set_cell_text(table.cell(3, 0), '클래스 타입', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(3, 0), YELLOW_SHADING)
    set_cell_text(table.cell(3, 1), class_type, font_size=9)
    set_cell_text(table.cell(3, 2), '', font_size=9)
    set_cell_text(table.cell(3, 3), '', font_size=9)
    set_cell_border_nil(table.cell(3, 1), ['right'])
    set_cell_border_nil(table.cell(3, 2), ['left', 'right'])
    set_cell_border_nil(table.cell(3, 3), ['left'])

    # Row 4
    cell0 = table.cell(4, 0)
    set_cell_shading(cell0, YELLOW_SHADING)
    set_cell_multiline(cell0, ['', '', '클래스 개요'], bold=True, font_size=9)
    for p in cell0.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    merge_cells_horizontal(table, 4, 1, 3)
    overview_lines = class_overview.split('\n')
    set_cell_multiline(table.cell(4, 1), overview_lines, font_size=9)

    # Row 5
    merge_cells_horizontal(table, 5, 0, 3)
    cell = table.cell(5, 0)
    set_cell_text(cell, '속성(Attributes)', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 6
    merge_cells_horizontal(table, 6, 0, 3)
    attr_lines = attributes_text.split('\n')
    set_cell_multiline(table.cell(6, 0), attr_lines, font_size=9)

    if has_ops:
        # Row 7
        merge_cells_horizontal(table, 7, 0, 3)
        cell = table.cell(7, 0)
        set_cell_text(cell, '오퍼레이션(Operations)', bold=True, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, YELLOW_SHADING)

        # Row 8
        merge_cells_horizontal(table, 8, 0, 3)
        ops_lines = operations_text.split('\n')
        set_cell_multiline(table.cell(8, 0), ops_lines, font_size=9)

    doc.add_paragraph()  # spacing


def make_overview(line1, line2):
    return '\n'.join([
        '개요',
        f'-    {line1}',
        '',
        '상세내용',
        f'-    {line2}',
    ])


# ============================================================
# DATA DEFINITIONS
# ============================================================

def build_all_classes():
    """Return list of (group_name, class_type_label, class_data_list) tuples."""
    groups = []

    # ================================================================
    # GROUP 1: Health
    # ================================================================
    health_classes = []

    health_classes.append(('Controller', [{
        'class_name': 'HealthController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'Health Controller를 정의한다.',
            'DMZ Agent 기동 상태, 등록된 Agent 목록(RCV/Loader/SND), 실행 중 Agent 반환 및 디버그용 DataSource 정보 조회 기능을 제공한다.'
        ),
        'attributes': '\n'.join([
            '@Value("${agent.zone}")',
            'private String zone;',
            '',
            'private final PipelineRegistry pipelineRegistry;',
            'private final PipelineService pipelineService;',
            'private final SyncDataSourceService syncDataSourceService;',
        ]),
        'operations': '\n'.join([
            '// Agent 기동 상태 및 등록 정보 조회',
            '@GetMapping("/health")',
            'public ResponseEntity<Map<String, Object>> health();',
            '',
            '// 디버그용: 캐시된 DataSource 설정 확인',
            '@GetMapping("/debug/datasources")',
            'public ResponseEntity<Map<String, Object>> debugDatasources();',
        ]),
    }]))

    groups.append(('Health', health_classes))

    # ================================================================
    # GROUP 2: Pipeline
    # ================================================================
    pipeline_classes = []

    # --- PipelineController ---
    pipeline_classes.append(('Controller', [{
        'class_name': 'PipelineController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'Pipeline Controller를 정의한다.',
            'Orchestrator로부터 실행 트리거를 받아 agentCode 기반 RCV/Loader/SND 파이프라인을 라우팅하고, 재동기화, 실행 상태/결과 조회, 테이블 정보 조회 기능을 제공한다.'
        ),
        'attributes': '\n'.join([
            'private final PipelineService pipelineService;',
            'private final PipelineRegistry pipelineRegistry;',
            'private final AgentConfigLoader agentConfigLoader;',
            'private static final DateTimeFormatter FORMATTER;',
        ]),
        'operations': '\n'.join([
            '// 파이프라인 실행 트리거 (비동기)',
            '@PostMapping("/execute")',
            'public ResponseEntity<Map<String, Object>> execute(@RequestBody Map<String, Object> request);',
            '',
            '// 재동기화 (관리자가 기간 지정)',
            '@PostMapping("/resync")',
            'public ResponseEntity<Map<String, Object>> resync(@RequestBody Map<String, Object> request);',
            '',
            '// 실행 상태 조회',
            '@GetMapping("/status/{executionId}")',
            'public ResponseEntity<?> getStatus(@PathVariable String executionId);',
            '',
            '// 실행 결과 상세 조회',
            '@GetMapping("/execution/{executionId}")',
            'public ResponseEntity<Map<String, Object>> getExecutionResult(@PathVariable String executionId);',
            '',
            '// WHERE 조건 대상 테이블 목록 조회',
            '@GetMapping("/{agentCode}/select-tables")',
            'public ResponseEntity<List<String>> getSelectTables(@PathVariable String agentCode);',
            '',
            '// 파이프라인 테이블 정보 조회',
            '@GetMapping("/{agentCode}/tables")',
            'public ResponseEntity<Map<String, Object>> getPipelineTables(@PathVariable String agentCode);',
        ]),
    }]))

    # --- PipelineService ---
    pipeline_classes.append(('Service', [{
        'class_name': 'PipelineService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'Pipeline Service를 정의한다.',
            'agentCode 기반 파이프라인 비동기 실행, OrchestratorClient 생성, DataSource 해석(Proxy 경유), 실행 결과 관리 및 Orchestrator 콜백 알림을 담당한다.'
        ),
        'attributes': '\n'.join([
            'private final PipelineRegistry pipelineRegistry;',
            'private final ExecutionService executionService;',
            'private final SyncDataSourceService syncDataSourceService;',
            '',
            '@Value("${agent.orchestrator-url}")',
            'private String orchestratorUrl;',
            '',
            '@Value("${agent.zone}")',
            'private String agentZone;',
            '',
            'private final Map<String, PipelineResult> executionResults;',
            'private final Set<String> runningAgentCodes;',
        ]),
        'operations': '\n'.join([
            '// 파이프라인 비동기 실행 (Orchestrator에서 호출)',
            '@Async("pipelineExecutor")',
            'public void executeAsync(String executionId, Map<String, Object> params);',
            '',
            '// 실행 결과 조회',
            'public PipelineResult getExecutionResult(String executionId);',
            '',
            '// 실행 중인 Agent 코드 목록 조회',
            'public Set<String> getRunningAgentCodes();',
        ]),
    },
    {
        'class_name': 'TargetRepositoryService',
        'class_type': 'Class <<Service>>',
        'overview': make_overview(
            'Target Repository Service를 정의한다.',
            'Target DB CRUD용 Repository Service로, 동적 EntityManager/JdbcTemplate을 사용하여 제원(SecJewon), 관측데이터(SecObsvdata), Link(LinkNgwis), IF_RSV 데이터의 조회/저장/삭제 및 JDBC batch UPSERT, IF→Target 변환을 담당한다.'
        ),
        'attributes': '\n'.join([
            'private final DynamicEntityManagerService entityManagerService;',
            'private static final int JDBC_BATCH_SIZE = 1000;',
        ]),
        'operations': '\n'.join([
            '// === SecJewon 조회 ===',
            '// 제원 목록 조회 (페이징, 검색, 정렬)',
            'public List<SecJewon> findJewonWithPaging(int offset, int limit, String search, String searchColumn, String sortColumn, String sortDirection);',
            '',
            '// 제원 총 개수 조회 (검색)',
            'public long countJewon(String search, String searchColumn);',
            '',
            '// === SecJewon 저장/삭제 ===',
            '// 제원 저장 (INSERT or UPDATE)',
            'public SecJewon saveJewon(SecJewon jewon);',
            '',
            '// 제원 일괄 저장',
            'public int saveAllJewon(List<SecJewon> jewonList);',
            '',
            '// 제원 삭제 (obsvCode 목록)',
            'public int deleteJewonByObsvCodes(List<String> obsvCodes);',
            '',
            '// === SecObsvdata ===',
            '// 관측데이터 목록 조회 (페이징, 검색, 정렬)',
            'public List<SecObsvdata> findObsvDataWithPaging(int offset, int limit, String search, String searchColumn, String sortColumn, String sortDirection);',
            '',
            '// 관측데이터 총 개수 조회 (검색)',
            'public long countObsvData(String search, String searchColumn);',
            '',
            '// 관측데이터 일괄 저장',
            'public int saveAllObsvData(List<SecObsvdata> obsvDataList);',
            '',
            '// === LinkNgwis ===',
            '// Link 조회',
            'public Optional<LinkNgwis> findLinkByObsvCode(String obsvCode);',
            '',
            '// Link 저장 (INSERT or UPDATE)',
            'public LinkNgwis saveLink(LinkNgwis link);',
            '',
            '// Link 업데이트 (마지막 동기화 시점 기록)',
            'public void updateLinkLastSync(String obsvCode, Date obsvDate, Time obsvTime);',
            '',
            '// === JDBC Batch UPSERT (성능 최적화) ===',
            '// 제원 JDBC batch UPSERT (ON CONFLICT)',
            'public int batchUpsertJewon(List<SecJewon> jewonList);',
            '',
            '// 관측데이터 JDBC batch UPSERT (ON CONFLICT)',
            'public int batchUpsertObsvdata(List<SecObsvdata> obsvList);',
            '',
            '// Link 테이블 JDBC batch UPSERT',
            'public int batchUpsertLinks(List<LinkNgwis> linkList);',
            '',
            '// === IF -> Target 변환 헬퍼 ===',
            '// IfRsvSecJewon -> SecJewon 변환',
            'public static SecJewon convertToSecJewon(IfRsvSecJewon ifJewon, String executionId);',
            '',
            '// IfRsvSecObsvdata -> SecObsvdata 변환',
            'public static SecObsvdata convertToSecObsvdata(IfRsvSecObsvdata ifData, String executionId);',
            '',
            '// === executionId 기반 조회 ===',
            '// 제원 조회 (executionId 기반)',
            'public List<SecJewon> findJewonByExecutionId(String executionId);',
            '',
            '// 관측데이터 조회 (executionId 기반)',
            'public List<SecObsvdata> findObsvDataByExecutionId(String executionId);',
            '',
            '// === sourceRefs 기반 조회 ===',
            '// 제원 조회 (sourceRefs 목록)',
            'public List<SecJewon> findJewonBySourceRefs(List<String> sourceRefsList);',
            '',
            '// 관측데이터 조회 (sourceRefs 목록)',
            'public List<SecObsvdata> findObsvDataBySourceRefs(List<String> sourceRefsList);',
            '',
            '// === IF_RSV 조회 ===',
            '// IF_RSV 제원 - PENDING/RESYNC 상태 조회',
            'public List<IfRsvSecJewon> findIfRsvJewonPending(String executionId);',
            '',
            '// IF_RSV 관측데이터 - PENDING/RESYNC 상태 조회',
            'public List<IfRsvSecObsvdata> findIfRsvObsvdataPending(String executionId);',
            '',
            '// IF_RSV 제원 - 동적 조건으로 조회 (resync용)',
            'public List<IfRsvSecJewon> findIfRsvJewonForResync(List<ExecutionCondition> conditions);',
            '',
            '// IF_RSV 관측데이터 - 동적 조건으로 조회 (resync용)',
            'public List<IfRsvSecObsvdata> findIfRsvObsvdataForResync(List<ExecutionCondition> conditions);',
        ]),
    }]))

    # --- Entities ---
    entity_list = []

    # SecJewon
    entity_list.append({
        'class_name': 'SecJewon',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'Target DB 제원(관측소 기본정보) 엔티티를 정의한다.',
            'Loader 파이프라인에서 IF_RSV의 데이터를 UPSERT하는 최종 목적지이다. source_refs로 원본 Source를 추적하며, execution_id로 실행 이력을 조회할 수 있다. 테이블: sec_jewon (UK: source_refs, IDX: execution_id)'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            '@Column(name = "id")',
            'private Long id;',
            '',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Column(name = "obsv_name")',
            'private String obsvName;',
            '',
            '@Column(name = "well")',
            'private Integer well;',
            '',
            '@Column(name = "sido")',
            'private String sido;',
            '',
            '@Column(name = "sigungu")',
            'private String sigungu;',
            '',
            '@Column(name = "upmyundo")',
            'private String upmyundo;',
            '',
            '@Column(name = "bunji")',
            'private String bunji;',
            '',
            '@Column(name = "ri")',
            'private String ri;',
            '',
            '@Column(name = "x")',
            'private String x;',
            '',
            '@Column(name = "y")',
            'private String y;',
            '',
            '@Column(name = "pyogo")',
            'private Double pyogo;',
            '',
            '@Column(name = "insdate")',
            'private Date insdate;',
            '',
            '@Column(name = "guldep")',
            'private Double guldep;',
            '',
            '@Column(name = "guldia")',
            'private Double guldia;',
            '',
            '@Column(name = "regdate")',
            'private Date regdate;',
            '',
            '@Column(name = "casing_height")',
            'private Double casingHeight;',
            '',
            '@Column(name = "source_refs", columnDefinition = "TEXT")',
            'private String sourceRefs;',
            '',
            '@Column(name = "link_status", length = 20)',
            'private String linkStatus;',
            '',
            '@Column(name = "execution_id", length = 100)',
            'private String executionId;',
        ]),
    })

    # SecObsvdata
    entity_list.append({
        'class_name': 'SecObsvdata',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'Target DB 관측데이터 엔티티를 정의한다.',
            'Loader 파이프라인에서 IF_RSV의 관측 시계열 데이터를 UPSERT하는 최종 목적지이다. 수위(gwdep), 수온(gwtemp), 전기전도도(ec) 등을 저장하며, source_refs로 원본 추적, execution_id로 실행 이력 조회가 가능하다. 테이블: sec_obsvdata (UK: source_refs, IDX: execution_id)'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            '@Column(name = "id")',
            'private Integer id;',
            '',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Column(name = "obsv_date")',
            'private Date obsvDate;',
            '',
            '@Column(name = "obsv_time")',
            'private Time obsvTime;',
            '',
            '@Column(name = "gwdep")',
            'private Double gwdep;',
            '',
            '@Column(name = "gwtemp")',
            'private Double gwtemp;',
            '',
            '@Column(name = "ec")',
            'private Double ec;',
            '',
            '@Column(name = "remark")',
            'private String remark;',
            '',
            '@Column(name = "source_refs", columnDefinition = "TEXT")',
            'private String sourceRefs;',
            '',
            '@Column(name = "link_status", length = 20)',
            'private String linkStatus;',
            '',
            '@Column(name = "execution_id", length = 100)',
            'private String executionId;',
        ]),
    })

    # LinkNgwis
    entity_list.append({
        'class_name': 'LinkNgwis',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            '동기화 시점 추적 엔티티를 정의한다.',
            '외부 업체 연계 시 관측소별 마지막 동기화 시점(obsv_date, obsv_time)을 기록하여 증분 동기화에 활용한다. 테이블: link_ngwis (PK: obsv_code)'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Column(name = "obsv_date")',
            'private LocalDateTime obsvDate;',
            '',
            '@Column(name = "obsv_time")',
            'private String obsvTime;',
            '',
            '@Column(name = "update_time")',
            'private LocalDateTime updateTime;',
        ]),
    })

    # SecJewonView
    entity_list.append({
        'class_name': 'SecJewonView',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            '외부 소스 DB의 제원(관측소 기본정보) 뷰 엔티티를 정의한다.',
            'RCV 파이프라인에서 외부 업체 DB의 제원 데이터를 SELECT할 때 사용되는 읽기 전용 뷰 엔티티이다. PK는 obsv_code이며, INSERT/UPDATE 대상이 아니다. 테이블(뷰): sec_jewon_view'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Column(name = "obsv_name")',
            'private String obsvName;',
            '',
            '@Column(name = "well")',
            'private Integer well;',
            '',
            '@Column(name = "sido")',
            'private String sido;',
            '',
            '@Column(name = "sigungu")',
            'private String sigungu;',
            '',
            '@Column(name = "upmyundo")',
            'private String upmyundo;',
            '',
            '@Column(name = "bunji")',
            'private String bunji;',
            '',
            '@Column(name = "ri")',
            'private String ri;',
            '',
            '@Column(name = "x")',
            'private String x;',
            '',
            '@Column(name = "y")',
            'private String y;',
            '',
            '@Column(name = "pyogo")',
            'private Double pyogo;',
            '',
            '@Column(name = "insdate")',
            'private Date insdate;',
            '',
            '@Column(name = "guldep")',
            'private Double guldep;',
            '',
            '@Column(name = "guldia")',
            'private Double guldia;',
            '',
            '@Column(name = "regdate")',
            'private Date regdate;',
            '',
            '@Column(name = "casing_height")',
            'private Double casingHeight;',
        ]),
    })

    # SecObsvdataView
    entity_list.append({
        'class_name': 'SecObsvdataView',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            '외부 소스 DB의 관측데이터 뷰 엔티티를 정의한다.',
            'RCV 파이프라인에서 외부 업체 DB의 관측 시계열 데이터를 SELECT할 때 사용되는 읽기 전용 뷰 엔티티이다. 복합PK: obsv_code + obsv_date + obsv_time. 테이블(뷰): sec_obsvdata_view'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Id',
            '@Column(name = "obsv_date")',
            'private Date obsvDate;',
            '',
            '@Id',
            '@Column(name = "obsv_time")',
            'private Time obsvTime;',
            '',
            '@Column(name = "gwdep")',
            'private Double gwdep;',
            '',
            '@Column(name = "gwtemp")',
            'private Double gwtemp;',
            '',
            '@Column(name = "ec")',
            'private Double ec;',
            '',
            '@Column(name = "remark")',
            'private String remark;',
        ]),
    })

    # SecObsvdataViewId
    entity_list.append({
        'class_name': 'SecObsvdataViewId',
        'class_type': 'Class <<IdClass>>',
        'overview': make_overview(
            'SecObsvdataView 복합 PK 클래스를 정의한다.',
            'SecObsvdataView 엔티티의 복합 기본키(obsv_code + obsv_date + obsv_time)를 표현하는 Serializable 클래스이다.'
        ),
        'attributes': '\n'.join([
            'private String obsvCode;',
            '',
            'private Date obsvDate;',
            '',
            'private Time obsvTime;',
        ]),
    })

    # IfRsvSecJewon
    entity_list.append({
        'class_name': 'IfRsvSecJewon',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'IF_RSV 수신용 제원 엔티티를 정의한다.',
            'RCV 파이프라인(relay-dmz-rsv-bojo)에서 외부 DB의 제원 데이터를 적재하고, Loader가 읽어서 Target(sec_jewon)으로 매핑하는 중간 테이블이다. PK: Auto-generated, UK: source_refs. 테이블: if_rsv_sec_jewon'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            '@Column(name = "id")',
            'private Integer id;',
            '',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Column(name = "obsv_name")',
            'private String obsvName;',
            '',
            '@Column(name = "well")',
            'private Integer well;',
            '',
            '@Column(name = "sido")',
            'private String sido;',
            '',
            '@Column(name = "sigungu")',
            'private String sigungu;',
            '',
            '@Column(name = "upmyundo")',
            'private String upmyundo;',
            '',
            '@Column(name = "bunji")',
            'private String bunji;',
            '',
            '@Column(name = "ri")',
            'private String ri;',
            '',
            '@Column(name = "x")',
            'private String x;',
            '',
            '@Column(name = "y")',
            'private String y;',
            '',
            '@Column(name = "pyogo")',
            'private Double pyogo;',
            '',
            '@Column(name = "insdate")',
            'private Date insdate;',
            '',
            '@Column(name = "guldep")',
            'private Double guldep;',
            '',
            '@Column(name = "guldia")',
            'private Double guldia;',
            '',
            '@Column(name = "regdate")',
            'private Date regdate;',
            '',
            '@Column(name = "casing_height")',
            'private Double casingHeight;',
            '',
            '@Column(name = "source_refs", columnDefinition = "TEXT")',
            'private String sourceRefs;',
            '',
            '@Builder.Default',
            '@Column(name = "link_status", length = 20)',
            'private String linkStatus = "PENDING";',
            '',
            '@Column(name = "extracted_at")',
            'private LocalDateTime extractedAt;',
            '',
            '@Column(name = "updated_at")',
            'private LocalDateTime updatedAt;',
            '',
            '@Column(name = "execution_id")',
            'private String executionId;',
        ]),
    })

    # IfRsvSecObsvdata
    entity_list.append({
        'class_name': 'IfRsvSecObsvdata',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'IF_RSV 수신용 관측데이터 엔티티를 정의한다.',
            'RCV 파이프라인에서 외부 DB의 관측 시계열 데이터를 적재하고, Loader가 읽어서 Target(sec_obsvdata)으로 매핑하는 중간 테이블이다. PK: Auto-generated, UK: source_refs. 테이블: if_rsv_sec_obsvdata'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@GeneratedValue(strategy = GenerationType.IDENTITY)',
            '@Column(name = "id")',
            'private Integer id;',
            '',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Column(name = "obsv_date")',
            'private Date obsvDate;',
            '',
            '@Column(name = "obsv_time")',
            'private Time obsvTime;',
            '',
            '@Column(name = "gwdep")',
            'private Double gwdep;',
            '',
            '@Column(name = "gwtemp")',
            'private Double gwtemp;',
            '',
            '@Column(name = "ec")',
            'private Double ec;',
            '',
            '@Column(name = "remark")',
            'private String remark;',
            '',
            '@Column(name = "source_refs", columnDefinition = "TEXT")',
            'private String sourceRefs;',
            '',
            '@Builder.Default',
            '@Column(name = "link_status", length = 20)',
            'private String linkStatus = "PENDING";',
            '',
            '@Column(name = "extracted_at")',
            'private LocalDateTime extractedAt;',
            '',
            '@Column(name = "updated_at")',
            'private LocalDateTime updatedAt;',
            '',
            '@Column(name = "execution_id")',
            'private String executionId;',
        ]),
    })

    # IfSndSecJewon
    entity_list.append({
        'class_name': 'IfSndSecJewon',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'IF_SND 송신용 제원 엔티티를 정의한다.',
            'SND 파이프라인에서 Target DB의 제원 데이터를 IF_SND 테이블로 추출할 때 사용된다. source_refs에 UK, execution_id에 인덱스가 설정되어 있다. 테이블: if_snd_sec_jewon'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@Column(name = "id")',
            'private Integer id;',
            '',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Column(name = "obsv_name")',
            'private String obsvName;',
            '',
            '@Column(name = "well")',
            'private Integer well;',
            '',
            '@Column(name = "sido")',
            'private String sido;',
            '',
            '@Column(name = "sigungu")',
            'private String sigungu;',
            '',
            '@Column(name = "upmyundo")',
            'private String upmyundo;',
            '',
            '@Column(name = "bunji")',
            'private String bunji;',
            '',
            '@Column(name = "ri")',
            'private String ri;',
            '',
            '@Column(name = "x")',
            'private String x;',
            '',
            '@Column(name = "y")',
            'private String y;',
            '',
            '@Column(name = "pyogo")',
            'private Double pyogo;',
            '',
            '@Column(name = "insdate")',
            'private Date insdate;',
            '',
            '@Column(name = "guldep")',
            'private Double guldep;',
            '',
            '@Column(name = "guldia")',
            'private Double guldia;',
            '',
            '@Column(name = "regdate")',
            'private Date regdate;',
            '',
            '@Column(name = "casing_height")',
            'private Double casingHeight;',
            '',
            '@Column(name = "source_refs", columnDefinition = "TEXT")',
            'private String sourceRefs;',
            '',
            '@Builder.Default',
            '@Column(name = "link_status", length = 20)',
            'private String linkStatus = "PENDING";',
            '',
            '@Column(name = "extracted_at")',
            'private LocalDateTime extractedAt;',
            '',
            '@Column(name = "updated_at")',
            'private LocalDateTime updatedAt;',
            '',
            '@Column(name = "execution_id")',
            'private String executionId;',
        ]),
    })

    # IfSndSecObsvdata
    entity_list.append({
        'class_name': 'IfSndSecObsvdata',
        'class_type': 'Class <<Entity>>',
        'overview': make_overview(
            'IF_SND 송신용 관측데이터 엔티티를 정의한다.',
            'SND 파이프라인에서 Target DB의 관측데이터를 IF_SND 테이블로 추출할 때 사용된다. 수위(gwdep), 수온(gwtemp), 전기전도도(ec) 등의 시계열 관측값을 보관한다. 테이블: if_snd_sec_obsvdata'
        ),
        'attributes': '\n'.join([
            '@Id',
            '@Column(name = "id")',
            'private Integer id;',
            '',
            '@Column(name = "obsv_code")',
            'private String obsvCode;',
            '',
            '@Column(name = "obsv_date")',
            'private Date obsvDate;',
            '',
            '@Column(name = "obsv_time")',
            'private Time obsvTime;',
            '',
            '@Column(name = "gwdep")',
            'private Double gwdep;',
            '',
            '@Column(name = "gwtemp")',
            'private Double gwtemp;',
            '',
            '@Column(name = "ec")',
            'private Double ec;',
            '',
            '@Column(name = "remark")',
            'private String remark;',
            '',
            '@Column(name = "source_refs", columnDefinition = "TEXT")',
            'private String sourceRefs;',
            '',
            '@Builder.Default',
            '@Column(name = "link_status", length = 20)',
            'private String linkStatus = "PENDING";',
            '',
            '@Column(name = "extracted_at")',
            'private LocalDateTime extractedAt;',
            '',
            '@Column(name = "updated_at")',
            'private LocalDateTime updatedAt;',
            '',
            '@Column(name = "execution_id")',
            'private String executionId;',
        ]),
    })

    pipeline_classes.append(('Entity', entity_list))

    groups.append(('Pipeline', pipeline_classes))

    return groups


# ============================================================
# DOCUMENT GENERATION
# ============================================================

def main():
    doc = Document()

    # Page margins 2cm
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('sync-agent-bojo 클래스명세서')
    run.font.size = Pt(16)
    run.font.name = '맑은 고딕'
    run.bold = True
    rPr = run._r.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>'))

    doc.add_paragraph()

    groups = build_all_classes()

    # Heading numbering
    # group_num tracks major number (1, 2, ...)
    group_num = 0

    for group_name, type_groups in groups:
        group_num += 1

        # Group heading (e.g. "1    Health")
        h1 = doc.add_heading(f'{group_num}    {group_name}', level=1)

        type_num = 0
        for type_label, class_list in type_groups:
            type_num += 1

            # Type heading (e.g. "1.1    Controller")
            h2 = doc.add_heading(f'{group_num}.{type_num}    {type_label}', level=2)

            for cls_idx, cls_data in enumerate(class_list, start=1):
                # Class heading (e.g. "1.1.1    HealthController")
                h3 = doc.add_heading(
                    f'{group_num}.{type_num}.{cls_idx}    {cls_data["class_name"]}',
                    level=3
                )

                # Package name based on type_label
                if type_label == 'Controller':
                    pkg = 'kr.go.gims.agent.bojo.controller'
                elif type_label == 'Service':
                    if cls_data['class_name'] == 'TargetRepositoryService':
                        pkg = 'kr.go.gims.agent.bojo.loader.repository'
                    else:
                        pkg = 'kr.go.gims.agent.bojo.pipeline'
                elif type_label == 'Entity':
                    name = cls_data['class_name']
                    if name.startswith('IfRsv'):
                        pkg = 'kr.go.gims.agent.bojo.entity.iftable.rsv'
                    elif name.startswith('IfSnd'):
                        pkg = 'kr.go.gims.agent.bojo.entity.iftable.snd'
                    elif name.startswith('Sec') and 'View' in name:
                        pkg = 'kr.go.gims.agent.bojo.entity.source'
                    elif name == 'SecObsvdataViewId':
                        pkg = 'kr.go.gims.agent.bojo.entity.source'
                    elif name == 'LinkNgwis':
                        pkg = 'kr.go.gims.agent.bojo.entity.target'
                    else:
                        pkg = 'kr.go.gims.agent.bojo.entity.target'
                else:
                    pkg = 'kr.go.gims.agent.bojo'

                create_class_table(
                    doc,
                    package_name=pkg,
                    class_id='',
                    class_name=cls_data['class_name'],
                    screen_id='',
                    class_type=cls_data['class_type'],
                    class_overview=cls_data['overview'],
                    attributes_text=cls_data['attributes'],
                    operations_text=cls_data.get('operations'),
                )

    # Save
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'output')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'sync-agent-bojo_NGIS_D23_클래스명세서v0.1.docx')
    doc.save(output_path)
    print(f'Generated: {output_path}')


if __name__ == '__main__':
    main()
