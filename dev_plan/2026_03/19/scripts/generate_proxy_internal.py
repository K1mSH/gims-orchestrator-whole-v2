# -*- coding: utf-8 -*-
"""
sync-proxy-internal NGIS D23 클래스명세서 생성 스크립트
- 2개 도메인 그룹, 3개 클래스
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

YELLOW_SHADING = 'FFFF00'
COL_WIDTHS_DXA = [1800, 2701, 1529, 2902]
ROW_HEIGHTS_WITH_OPS = [311, 321, 297, 297, 1300, 261, 275, 261, 258]
ROW_HEIGHTS_NO_OPS = [311, 321, 297, 297, 1300, 261, 275]


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _make_run(p, text, bold=False, font_size=9):
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
    rPr.append(rFonts)
    return run


def set_cell_text(cell, text, bold=False, font_size=9, alignment=None):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _make_run(p, text, bold=bold, font_size=font_size)


def set_cell_multiline(cell, lines_list, bold=False, font_size=9):
    for p in cell.paragraphs:
        p.clear()
    for i, line in enumerate(lines_list):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _make_run(p, line, bold=bold, font_size=font_size)


def merge_cells_horizontal(table, row_idx, start_col, end_col):
    table.cell(row_idx, start_col).merge(table.cell(row_idx, end_col))


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))


def set_table_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing = tblPr.find(qn('w:tblLayout'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(layout)


def set_table_indent(table, indent_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    indent_xml = f'<w:tblInd {nsdecls("w")} w:w="{indent_dxa}" w:type="dxa"/>'
    existing = tblPr.find(qn('w:tblInd'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(indent_xml))


def set_row_height(row, height_dxa):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{height_dxa}" w:hRule="atLeast"/>')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def set_column_widths(table, widths_dxa):
    tbl = table._tbl
    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in widths_dxa:
        grid_xml += f'<w:gridCol w:w="{w}"/>'
    grid_xml += '</w:tblGrid>'
    tblPr = tbl.tblPr
    tblPr.addnext(parse_xml(grid_xml))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_dxa):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                tcPr.append(parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{widths_dxa[i]}" w:type="dxa"/>'
                ))


def set_cell_border_nil(cell, sides):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_parts = []
    for side in sides:
        borders_parts.append(f'<w:{side} w:val="nil"/>')
    borders_xml = f'<w:tcBorders {nsdecls("w")}>{"".join(borders_parts)}</w:tcBorders>'
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(borders_xml))


def create_class_table(doc, package_name, class_id, class_name, screen_id,
                       class_type, class_overview, attributes_text,
                       operations_text=None):
    """Create a single class specification table."""
    has_ops = operations_text is not None
    num_rows = 9 if has_ops else 7

    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table)
    set_table_layout_fixed(table)
    set_table_indent(table, 125)
    set_column_widths(table, COL_WIDTHS_DXA)

    heights = ROW_HEIGHTS_WITH_OPS if has_ops else ROW_HEIGHTS_NO_OPS
    for i, row in enumerate(table.rows):
        if i < len(heights):
            set_row_height(row, heights[i])

    # Row 0: Header
    merge_cells_horizontal(table, 0, 0, 3)
    cell = table.cell(0, 0)
    set_cell_text(cell, '클래스 명세서[설계]', bold=True, font_size=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 1: 패키지명 / 클래스ID
    set_cell_text(table.cell(1, 0), '패키지명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(1, 0), YELLOW_SHADING)
    set_cell_text(table.cell(1, 1), package_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 2), '클래스ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(1, 2), YELLOW_SHADING)
    set_cell_text(table.cell(1, 3), class_id, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 2: 클래스명 / 화면ID
    set_cell_text(table.cell(2, 0), '클래스명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(2, 0), YELLOW_SHADING)
    set_cell_text(table.cell(2, 1), class_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 2), '화면ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(2, 2), YELLOW_SHADING)
    set_cell_text(table.cell(2, 3), screen_id, font_size=9)

    # Row 3: 클래스 타입
    set_cell_text(table.cell(3, 0), '클래스 타입', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(3, 0), YELLOW_SHADING)
    set_cell_text(table.cell(3, 1), class_type, font_size=9)
    set_cell_text(table.cell(3, 2), '', font_size=9)
    set_cell_text(table.cell(3, 3), '', font_size=9)
    set_cell_border_nil(table.cell(3, 1), ['right'])
    set_cell_border_nil(table.cell(3, 2), ['left', 'right'])
    set_cell_border_nil(table.cell(3, 3), ['left'])

    # Row 4: 클래스 개요
    cell0 = table.cell(4, 0)
    set_cell_shading(cell0, YELLOW_SHADING)
    set_cell_multiline(cell0, ['', '', '클래스 개요'], bold=True, font_size=9)
    for p in cell0.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    merge_cells_horizontal(table, 4, 1, 3)
    overview_lines = class_overview.split('\n')
    set_cell_multiline(table.cell(4, 1), overview_lines, font_size=9)

    # Row 5: 속성(Attributes) header
    merge_cells_horizontal(table, 5, 0, 3)
    cell = table.cell(5, 0)
    set_cell_text(cell, '속성(Attributes)', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 6: attributes content
    merge_cells_horizontal(table, 6, 0, 3)
    attr_lines = attributes_text.split('\n')
    set_cell_multiline(table.cell(6, 0), attr_lines, font_size=9)

    if has_ops:
        # Row 7: 오퍼레이션(Operations) header
        merge_cells_horizontal(table, 7, 0, 3)
        cell = table.cell(7, 0)
        set_cell_text(cell, '오퍼레이션(Operations)', bold=True, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, YELLOW_SHADING)

        # Row 8: operations content
        merge_cells_horizontal(table, 8, 0, 3)
        ops_lines = operations_text.split('\n')
        set_cell_multiline(table.cell(8, 0), ops_lines, font_size=9)

    doc.add_paragraph()  # spacing


# ============================================================
# Overview helper
# ============================================================
def make_overview(line1, line2):
    return '\n'.join([
        '개요',
        f'-    {line1}',
        '',
        '상세내용',
        f'-    {line2}',
    ])


# ============================================================
# DATA DEFINITIONS
# ============================================================

PACKAGE_MAP = {
    'Health': 'kr.go.gims.proxy.internal.controller',
    'ConnectionInfo': 'kr.go.gims.proxy.internal',
}


def build_all_classes():
    """Return list of (group_name, class_type_label, class_data_list) tuples."""
    groups = []

    # ================================================================
    # GROUP 1: Health
    # ================================================================
    health_classes = []

    health_classes.append(('Controller', [{
        'class_name': 'HealthController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'Health Controller를 정의한다.',
            '내부망 DB 프록시 전용 헬스체크 및 캐시된 DataSource 디버그 정보를 제공한다.'
        ),
        'attributes': '\n'.join([
            '@Value("${agent.zone}")',
            'private String zone;',
            'private final ProxyDataSourceService proxyDataSourceService;',
        ]),
        'operations': '\n'.join([
            '// 헬스체크 (status, appName, type, zone 반환)',
            '@GetMapping("/health")',
            'public ResponseEntity<Map<String, Object>> health();',
            '',
            '// 캐시된 DataSource 디버그 정보 조회',
            '@GetMapping("/debug/datasources")',
            'public ResponseEntity<Map<String, Object>> debugDatasources();',
        ]),
    }]))

    groups.append(('Health', health_classes))

    # ================================================================
    # GROUP 2: ConnectionInfo
    # ================================================================
    conn_classes = []

    # --- ConnectionInfoController ---
    conn_classes.append(('Controller', [{
        'class_name': 'ConnectionInfoController',
        'class_type': 'Class <<Controller>>',
        'overview': make_overview(
            'ConnectionInfo Controller를 정의한다.',
            'Agent용 connection-info 프록시 엔드포인트. Agent가 파이프라인 실행 시 DB 자격증명을 얻기 위해 호출하며, Orchestrator의 connection-info API를 패스스루한다.'
        ),
        'attributes': '\n'.join([
            '@Value("${agent.orchestrator-url:http://localhost:8080}")',
            'private String orchestratorUrl;',
        ]),
        'operations': '\n'.join([
            '// datasource 연결 정보 패스스루 조회 (Orchestrator → Agent)',
            '@GetMapping("/{datasourceId}/connection-info")',
            'public ResponseEntity<Map<String, Object>> getConnectionInfo(@PathVariable String datasourceId);',
        ]),
    }]))

    # --- ProxyDataSourceService ---
    conn_classes.append(('Service', [{
        'class_name': 'ProxyDataSourceService',
        'class_type': 'Class <<Service>> implements DataSourceProvider',
        'overview': make_overview(
            'Proxy DataSource 서비스를 정의한다.',
            'DB 프록시 전용 DataSourceProvider 구현체. 메모리 캐시 → Orchestrator API 조회 → Spring 기본 DataSource 순서의 Fallback 전략으로 JdbcTemplate을 제공한다.'
        ),
        'attributes': '\n'.join([
            'private final JdbcTemplate defaultJdbcTemplate;',
            'private final PasswordEncryptor passwordEncryptor;',
            '@Value("${agent.orchestrator-url:http://localhost:8080}")',
            'private String orchestratorUrl;',
            'private final Map<String, DataSourceInfo> cachedDataSourceInfos;',
            'private final Map<String, HikariDataSource> dataSources;',
            'private final Map<String, JdbcTemplate> jdbcTemplates;',
        ]),
        'operations': '\n'.join([
            '// Source datasourceId 조회',
            '@Override',
            'public String getSourceDatasourceId();',
            '',
            '// Target datasourceId 조회',
            '@Override',
            'public String getTargetDatasourceId();',
            '',
            '// Agent 타입 반환 ("PROXY")',
            '@Override',
            'public String getAgentType();',
            '',
            '// datasourceId별 DB 타입 조회',
            '@Override',
            'public String getDbType(String datasourceId);',
            '',
            '// datasourceId별 JdbcTemplate 조회 (캐시 → Orchestrator → 기본 DataSource)',
            '@Override',
            'public JdbcTemplate getJdbcTemplate(String datasourceId);',
            '',
            '// 캐시된 DataSourceInfo 맵 반환',
            'public Map<String, DataSourceInfo> getCachedDataSourceInfos();',
            '',
            '// 애플리케이션 종료 시 DataSource 정리',
            '@PreDestroy',
            'public void closeAll();',
        ]),
    }]))

    groups.append(('ConnectionInfo', conn_classes))

    return groups


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(9)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    doc.add_heading('sync-proxy-internal 클래스명세서', level=0)

    groups = build_all_classes()

    for group_idx, (group_name, type_list) in enumerate(groups, start=1):
        # Group heading (level 1)
        doc.add_heading(f'{group_idx}    {group_name}', level=1)

        sub_idx = 0
        for type_label, class_list in type_list:
            sub_idx += 1
            # Type heading (level 2)
            doc.add_heading(f'{group_idx}.{sub_idx}    {type_label}', level=2)

            for cls_idx, cls in enumerate(class_list, start=1):
                # Class heading (level 3)
                doc.add_heading(
                    f'{group_idx}.{sub_idx}.{cls_idx}    {cls["class_name"]}',
                    level=3
                )

                create_class_table(
                    doc,
                    package_name=PACKAGE_MAP[group_name],
                    class_id='',
                    class_name=cls['class_name'],
                    screen_id='',
                    class_type=cls['class_type'],
                    class_overview=cls['overview'],
                    attributes_text=cls['attributes'],
                    operations_text=cls.get('operations'),
                )

    # Save
    output_dir = r'D:\dev\claude\GIMS\orchestrator_v2\dev_plan\2026_03\19\output'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(
        output_dir, 'sync-proxy-internal_NGIS_D23_클래스명세서v0.1.docx'
    )
    doc.save(output_path)
    print(f'클래스명세서 생성 완료: {output_path}')
    print(f'파일 크기: {os.path.getsize(output_path)} bytes')

    # Count classes for verification
    total = sum(len(cl) for _, tl in groups for _, cl in tl)
    print(f'총 클래스 수: {total}개')


if __name__ == '__main__':
    main()
