# -*- coding: utf-8 -*-
"""
Entity 클래스명세서 샘플 생성 스크립트
- Agent 엔티티 1개로 양식 테스트
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

YELLOW_SHADING = 'FFFF00'
COL_WIDTHS_DXA = [1800, 2701, 1529, 2902]
ROW_HEIGHTS_NO_OPS = [311, 321, 297, 297, 1300, 261, 275]


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _make_run(p, text, bold=False, font_size=9):
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
    rPr.append(rFonts)
    return run


def set_cell_text(cell, text, bold=False, font_size=9, alignment=None):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _make_run(p, text, bold=bold, font_size=font_size)


def set_cell_multiline(cell, lines_list, bold=False, font_size=9):
    for p in cell.paragraphs:
        p.clear()
    for i, line in enumerate(lines_list):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _make_run(p, line, bold=bold, font_size=font_size)


def merge_cells_horizontal(table, row_idx, start_col, end_col):
    table.cell(row_idx, start_col).merge(table.cell(row_idx, end_col))


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))


def set_table_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing = tblPr.find(qn('w:tblLayout'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(layout)


def set_table_indent(table, indent_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    indent_xml = f'<w:tblInd {nsdecls("w")} w:w="{indent_dxa}" w:type="dxa"/>'
    existing = tblPr.find(qn('w:tblInd'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(indent_xml))


def set_row_height(row, height_dxa):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{height_dxa}" w:hRule="atLeast"/>')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def set_column_widths(table, widths_dxa):
    tbl = table._tbl
    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in widths_dxa:
        grid_xml += f'<w:gridCol w:w="{w}"/>'
    grid_xml += '</w:tblGrid>'
    tblPr = tbl.tblPr
    tblPr.addnext(parse_xml(grid_xml))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_dxa):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                tcPr.append(parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{widths_dxa[i]}" w:type="dxa"/>'
                ))


def set_cell_border_nil(cell, sides):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_parts = []
    for side in sides:
        borders_parts.append(f'<w:{side} w:val="nil"/>')
    borders_xml = f'<w:tcBorders {nsdecls("w")}>{"".join(borders_parts)}</w:tcBorders>'
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(borders_xml))


def create_class_table(doc, package_name, class_id, class_name, screen_id,
                       class_type, class_overview, attributes_text,
                       operations_text=None):
    """Create a single class specification table."""
    has_ops = operations_text is not None
    num_rows = 9 if has_ops else 7

    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table)
    set_table_layout_fixed(table)
    set_table_indent(table, 125)
    set_column_widths(table, COL_WIDTHS_DXA)

    heights = [311, 321, 297, 297, 1300, 261, 275, 261, 258] if has_ops else ROW_HEIGHTS_NO_OPS
    for i, row in enumerate(table.rows):
        if i < len(heights):
            set_row_height(row, heights[i])

    # Row 0: Header
    merge_cells_horizontal(table, 0, 0, 3)
    cell = table.cell(0, 0)
    set_cell_text(cell, '클래스 명세서[설계]', bold=True, font_size=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 1: 패키지명 / 클래스ID
    set_cell_text(table.cell(1, 0), '패키지명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(1, 0), YELLOW_SHADING)
    set_cell_text(table.cell(1, 1), package_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 2), '클래스ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(1, 2), YELLOW_SHADING)
    set_cell_text(table.cell(1, 3), class_id, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 2: 클래스명 / 화면ID
    set_cell_text(table.cell(2, 0), '클래스명', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(2, 0), YELLOW_SHADING)
    set_cell_text(table.cell(2, 1), class_name, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 2), '화면ID', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(table.cell(2, 2), YELLOW_SHADING)
    set_cell_text(table.cell(2, 3), screen_id, font_size=9)

    # Row 3: 클래스 타입
    set_cell_text(table.cell(3, 0), '클래스 타입', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(3, 0), YELLOW_SHADING)
    set_cell_text(table.cell(3, 1), class_type, font_size=9)
    set_cell_text(table.cell(3, 2), '', font_size=9)
    set_cell_text(table.cell(3, 3), '', font_size=9)
    set_cell_border_nil(table.cell(3, 1), ['right'])
    set_cell_border_nil(table.cell(3, 2), ['left', 'right'])
    set_cell_border_nil(table.cell(3, 3), ['left'])

    # Row 4: 클래스 개요
    cell0 = table.cell(4, 0)
    set_cell_shading(cell0, YELLOW_SHADING)
    overview_cell0_lines = ['', '', '클래스 개요']
    set_cell_multiline(cell0, overview_cell0_lines, bold=True, font_size=9)
    for p in cell0.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    merge_cells_horizontal(table, 4, 1, 3)
    overview_lines = class_overview.split('\n')
    set_cell_multiline(table.cell(4, 1), overview_lines, font_size=9)

    # Row 5: 속성(Attributes) header
    merge_cells_horizontal(table, 5, 0, 3)
    cell = table.cell(5, 0)
    set_cell_text(cell, '속성(Attributes)', bold=True, font_size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, YELLOW_SHADING)

    # Row 6: attributes content
    merge_cells_horizontal(table, 6, 0, 3)
    attr_lines = attributes_text.split('\n')
    set_cell_multiline(table.cell(6, 0), attr_lines, font_size=9)

    if has_ops:
        # Row 7: 오퍼레이션(Operations) header
        merge_cells_horizontal(table, 7, 0, 3)
        cell = table.cell(7, 0)
        set_cell_text(cell, '오퍼레이션(Operations)', bold=True, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, YELLOW_SHADING)

        # Row 8: operations content
        merge_cells_horizontal(table, 8, 0, 3)
        ops_lines = operations_text.split('\n')
        set_cell_multiline(table.cell(8, 0), ops_lines, font_size=9)

    doc.add_paragraph()  # spacing


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(9)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    doc.add_heading('sync-orchestrator 클래스명세서 (Entity 샘플)', level=0)

    # ── Heading 구조 ──
    doc.add_heading('1    Agent', level=1)
    doc.add_heading('1.5    Entity', level=2)
    doc.add_heading('1.5.1    Agent', level=3)

    # ── Agent Entity 클래스 명세서 ──
    attributes_text = '\n'.join([
        '@Id @GeneratedValue(strategy = GenerationType.IDENTITY)',
        'private Long id;                              // 에이전트 PK',
        '',
        '@Column(name = "agent_code", unique = true, nullable = false)',
        'private String agentCode;                     // 에이전트 고유 코드',
        '',
        '@Column(name = "agent_name", nullable = false)',
        'private String agentName;                     // 에이전트 표시명',
        '',
        '@Column(name = "endpoint_url", nullable = false)',
        'private String endpointUrl;                   // REST 엔드포인트 URL',
        '',
        '@Column(name = "zone", nullable = false)',
        'private String zone;                          // 네트워크 존 (DMZ/INTERNAL)',
        '',
        '@Column(name = "is_active")',
        'private Boolean isActive;                     // 활성화 여부',
        '',
        '@Enumerated(EnumType.STRING)',
        '@Column(name = "agent_type", nullable = false)',
        'private AgentType agentType;                  // 에이전트 유형 (RCV/LOADER/SND/DB_CON_PROXY)',
        '',
        '@Column(name = "datasource_tag")',
        'private String datasourceTag;                 // 데이터소스 태그',
        '',
        '@Column(name = "source_datasource_id")',
        'private String sourceDatasourceId;            // 소스 데이터소스 ID',
        '',
        '@Column(name = "target_datasource_id")',
        'private String targetDatasourceId;            // 타겟 데이터소스 ID',
        '',
        '@Column(name = "description", columnDefinition = "TEXT")',
        'private String description;                   // 설명',
        '',
        '@Enumerated(EnumType.STRING)',
        '@Column(name = "status")',
        'private AgentStatus status;                   // 상태 (ONLINE/OFFLINE/ERROR)',
        '',
        '@Column(name = "last_executed_at")',
        'private LocalDateTime lastExecutedAt;         // 마지막 실행 시각',
        '',
        '@Column(name = "last_execution_status")',
        'private String lastExecutionStatus;           // 마지막 실행 결과',
        '',
        '@CreationTimestamp',
        '@Column(name = "created_at", updatable = false)',
        'private LocalDateTime createdAt;              // 생성 시각',
        '',
        '@OneToMany(mappedBy = "agent", cascade = CascadeType.ALL, orphanRemoval = true)',
        'private List<Schedule> schedules;             // 스케줄 목록',
        '',
        '@OneToMany(mappedBy = "agent", cascade = CascadeType.ALL, orphanRemoval = true)',
        'private List<AgentChainMember> chainMembers;  // 체인 멤버 목록',
        '',
        '@OneToMany(mappedBy = "agent", cascade = CascadeType.ALL, orphanRemoval = true)',
        'private List<AgentTable> agentTables;         // Agent 테이블 목록',
        '',
        '@Column(name = "retention_config", columnDefinition = "TEXT")',
        'private String retentionConfig;               // 데이터 보존(Retention) 설정 JSON',
    ])

    overview_lines = '\n'.join([
        '개요',
        '-    Agent 엔티티를 정의한다.',
        '',
        '상세내용',
        '-    테이블: agent. 동기화 에이전트 정보를 관리하는 JPA 엔티티로,',
        '     에이전트 코드/이름/URL/존/유형/상태/실행이력/데이터소스 연결/Retention 설정 등을 포함한다.',
        '     Schedule, AgentChainMember, AgentTable과 1:N 연관관계를 가진다.',
    ])

    create_class_table(
        doc,
        package_name='com.sync.orchestrator.domain.agent',
        class_id='',
        class_name='Agent',
        screen_id='',
        class_type='Class <<Entity>>',
        class_overview=overview_lines,
        attributes_text=attributes_text,
        operations_text=None,  # Entity는 오퍼레이션 없음
    )

    # Save
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                              '' if os.path.basename(os.path.dirname(os.path.abspath(__file__))) == '19' else '')
    output_dir = r'D:\dev\claude\GIMS\orchestrator_v2\dev_plan\2026_03\19\output'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'entity_sample.docx')
    doc.save(output_path)
    print(f'Entity 샘플 생성 완료: {output_path}')
    print(f'파일 크기: {os.path.getsize(output_path)} bytes')


if __name__ == '__main__':
    main()
